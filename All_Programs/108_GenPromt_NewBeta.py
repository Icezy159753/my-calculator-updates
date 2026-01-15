# 108_GenPromt_v2.py - PyQt6 Version with Dark Mode
# GenPromt - Thai to English Translator for Survey Data
# =====================================================

import sys
import os
import time
import json
import threading
import webbrowser
import re

# PyQt6 imports
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QComboBox, QTextEdit, QFileDialog,
    QMessageBox, QGroupBox, QFrame, QProgressBar, QSplitter,
    QDialog, QTableWidget, QTableWidgetItem, QLineEdit, QHeaderView, QInputDialog
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QSize, QTimer
from PyQt6.QtGui import QFont, QIcon, QPalette, QColor

# External libraries
try:
    import requests
    from docx import Document
    import openpyxl
    import pyperclip
    import pyreadstat
    import pandas as pd
    import tempfile
    LIBS_INSTALLED = True
except ImportError:
    LIBS_INSTALLED = False

# =====================================================
# OpenRouter API key loading
# - Priority: env var OPENROUTER_API_KEY -> openrouter.json (exe/_internal/appdata)
# - Avoid hardcoding keys in source to prevent leaks in GitHub Actions
# =====================================================
OPENROUTER_API_KEY = None
OPENROUTER_CONFIG_FILENAME = "openrouter.json"


def _resolve_resource_path(filename):
    candidates = []
    if getattr(sys, "frozen", False):
        if hasattr(sys, "_MEIPASS"):
            candidates.append(os.path.join(sys._MEIPASS, filename))
        exe_dir = os.path.dirname(sys.executable)
        candidates.append(os.path.join(exe_dir, filename))
        candidates.append(os.path.join(exe_dir, "_internal", filename))
    module_dir = os.path.dirname(os.path.abspath(__file__))
    candidates.append(os.path.join(module_dir, filename))
    candidates.append(os.path.join(os.path.dirname(module_dir), filename))
    for path in candidates:
        if os.path.exists(path):
            return path
    return os.path.join(module_dir, filename)


def _get_appdata_dir():
    base_dir = os.environ.get("APPDATA")
    if not base_dir:
        base_dir = os.path.join(os.path.expanduser("~"), "AppData", "Roaming")
    return os.path.join(base_dir, "GenPromt")


def _get_openrouter_read_paths():
    paths = []
    if getattr(sys, "frozen", False):
        if hasattr(sys, "_MEIPASS"):
            paths.append(os.path.join(sys._MEIPASS, OPENROUTER_CONFIG_FILENAME))
        exe_dir = os.path.dirname(sys.executable)
        paths.append(os.path.join(exe_dir, OPENROUTER_CONFIG_FILENAME))
        paths.append(os.path.join(exe_dir, "_internal", OPENROUTER_CONFIG_FILENAME))
    paths.append(os.path.join(_get_appdata_dir(), OPENROUTER_CONFIG_FILENAME))
    module_dir = os.path.dirname(os.path.abspath(__file__))
    paths.append(os.path.join(module_dir, OPENROUTER_CONFIG_FILENAME))
    paths.append(os.path.join(os.path.dirname(module_dir), OPENROUTER_CONFIG_FILENAME))
    return paths


def _get_openrouter_write_paths():
    paths = []
    if getattr(sys, "frozen", False):
        exe_dir = os.path.dirname(sys.executable)
        paths.append(os.path.join(exe_dir, OPENROUTER_CONFIG_FILENAME))
        paths.append(os.path.join(exe_dir, "_internal", OPENROUTER_CONFIG_FILENAME))
    else:
        module_dir = os.path.dirname(os.path.abspath(__file__))
        paths.append(os.path.join(module_dir, OPENROUTER_CONFIG_FILENAME))
        paths.append(os.path.join(os.path.dirname(module_dir), OPENROUTER_CONFIG_FILENAME))
    paths.append(os.path.join(_get_appdata_dir(), OPENROUTER_CONFIG_FILENAME))
    return paths


def get_openrouter_api_key():
    global OPENROUTER_API_KEY
    if OPENROUTER_API_KEY is not None:
        return OPENROUTER_API_KEY

    env_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if env_key:
        OPENROUTER_API_KEY = env_key
        return OPENROUTER_API_KEY

    for config_path in _get_openrouter_read_paths():
        try:
            if not os.path.exists(config_path):
                continue
            with open(config_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            file_key = str(data.get("api_key", "")).strip()
            if file_key:
                OPENROUTER_API_KEY = file_key
                return OPENROUTER_API_KEY
        except Exception:
            continue

    OPENROUTER_API_KEY = ""
    return OPENROUTER_API_KEY


def save_openrouter_api_key(api_key):
    global OPENROUTER_API_KEY
    payload = json.dumps({"api_key": api_key})
    for config_path in _get_openrouter_write_paths():
        try:
            os.makedirs(os.path.dirname(config_path), exist_ok=True)
            with open(config_path, "w", encoding="utf-8") as f:
                f.write(payload)
            OPENROUTER_API_KEY = api_key
            return config_path
        except Exception:
            continue
    return ""

# --- Default Model ---
DEFAULT_MODEL = "google/gemini-3-flash-preview"

# --- Prompt Templates ---
PROMPT_JOD = """คุยไทยนะ
จากข้อมูลต่อไปนี้:

[ *** วางข้อมูล Word และ Excel *** ]

ให้ดำเนินการดังนี้:

## กฎสำคัญที่สุด (ห้ามละเมิด):
- **ห้ามแปลภาษาเอง** อย่างเด็ดขาด
- **ห้ามเขียน/สรุป/เรียบเรียงข้อความภาษาอังกฤษขึ้นมาใหม่** อย่างเด็ดขาด (ห้าม paraphrase)
- **ห้ามเอาข้อความภาษาอังกฤษจากข้ออื่นมาใส่** อย่างเด็ดขาด (เช่น ห้ามเอาอังกฤษของ S36 ไปใส่ให้ S37)
- **ต้องตรวจสอบเลขข้อให้ตรงกัน** ก่อนนำไปใช้
- ✅ **ข้อความอังกฤษที่ใส่ได้ ต้องเป็น "การคัดลอกตรงตัว (Exact Copy)" จากต้นฉบับเท่านั้น**
  - หมายถึง: ตัวอักษร/คำ/เครื่องหมาย ต้องตรงกับข้อความที่ปรากฏในข้อมูลต้นฉบับ
  - ห้ามแต่งเพิ่ม ห้ามสรุป ห้ามขยายความ ห้ามเติมคำเชื่อม
- ถ้าหาข้อความภาษาอังกฤษไม่เจอ หรือ ไม่แน่ใจว่าตรงกัน -> **ต้องเว้นว่างไว้เท่านั้น** (ปล่อยค่าว่างจริง ๆ ไม่ใส่อะไรเลย)

## วิธีหาข้อความภาษาอังกฤษในแบบสอบถาม (สำคัญมาก):
- แบบสอบถามเป็นตาราง 2 ภาษา โดยภาษาอังกฤษอาจ:
  1) อยู่บรรทัดถัดไป/ใต้ภาษาไทย
  2) อยู่ท้ายประโยคไทยใน "บล็อกเดียวกัน"
  3) อยู่คนละบรรทัดแต่ยังอยู่ใน "ส่วน/กรอบ/บล็อกของข้อเดียวกัน"
- **ต้องดูบล็อกที่มีเลขข้อเดียวกันเท่านั้น**
  - เช่น Name=s37 ต้องดึงอังกฤษจากบล็อก S37 เท่านั้น
  - ห้ามข้ามไปใช้ประโยคอังกฤษที่อยู่ในส่วนอื่น แม้จะใกล้กัน

## นิยาม "บล็อกของข้อเดียวกัน":
- เริ่มตั้งแต่บรรทัดที่พบเลขข้อ (เช่น "S37") จนก่อนถึงเลขข้อถัดไป (เช่น "S38" หรือ "S0" หรือ "Q1" ฯลฯ)
- หรือถ้าเป็นตาราง: อยู่ในกรอบ/แถว/ส่วนของเลขข้อนั้นเดียวกัน

## ขั้นตอนการทำงาน:

1) **ระบุคู่ข้อมูล:** สำหรับแต่ละแถวในตารางข้อมูลสุดท้าย (`Name`, `VAR_THA`):
   - หา "บล็อก" ที่มีเลขข้อตรงกับ `Name` (เช่น Name=s37 -> บล็อก S37)
   - ในบล็อกนั้น มองหาข้อความภาษาอังกฤษที่อยู่ใต้/ถัดจากภาษาไทย หรือท้ายบล็อก

2) **สร้างตารางผลลัพธ์:** สร้างตารางที่มี 2 คอลัมน์: `Name`, `VAR_ENG`

3) **เติมข้อมูล (เข้มงวด):**
   - นำ `Name` มาใส่เหมือนเดิม
   - **เฉพาะกรณีที่พบข้อความภาษาอังกฤษในบล็อกที่มีเลขข้อตรงกัน และคัดลอกได้ตรงตัวจากต้นฉบับเท่านั้น:**
     - ให้ใส่ `VAR_ENG` = "Sxx) " + (ข้อความอังกฤษที่คัดลอกตรงตัว)
   - **หากไม่เจอ หรือ ไม่แน่ใจว่าตรงกัน -> ต้องเว้นว่าง `VAR_ENG` ไว้ (ค่าว่างจริง ๆ)**
   - **ห้ามใส่ข้อความภาษาไทย** ลงใน `VAR_ENG`

4) ✅ **Self-check ก่อนส่งคำตอบ (ห้ามข้าม):**
   - ตรวจทุกค่า `VAR_ENG` ทีละบรรทัดว่า "ข้อความอังกฤษหลัง `Sxx)`" สามารถพบเป็นข้อความเดียวกันในข้อมูลต้นฉบับ (Exact substring) จริงหรือไม่
   - ถ้า **หาในต้นฉบับไม่พบแบบตรงตัว** -> ให้ลบค่านั้น แล้วเว้นว่าง

5) **จัดรูปแบบตาราง:**
   - ใช้ **Tab** เป็นตัวคั่นระหว่างคอลัมน์
   - ส่งกลับเป็น Table Markdown ที่มี 2 คอลัมน์: `Name`, `VAR_ENG`
   - จัดให้อยู่ในรูปแบบ Codebox

## ตัวอย่างย้ำ (ห้ามผิด):
- ถ้า A1 มีแต่ภาษาไทย (ไม่มีอังกฤษอยู่ในต้นฉบับ) -> ต้องเว้นว่างเท่านั้น ห้ามแปลเอง
- ถ้าบล็อก S37 มีอังกฤษว่า "Are you interested in participating in this research – HVT Phase?" -> ต้องคัดลอกประโยคนี้ตรงตัวเท่านั้น ห้ามสรุป/ขยาย
"""

PROMPT_CODE = """คุยไทยนะ
จากข้อมูลต่อไปนี้:

[ *** วางข้อมูล Word และ Excel *** ]

ให้ดำเนินการดังนี้:

## กฎสำคัญที่สุด (ห้ามละเมิด):
- **ห้ามแปลภาษาเอง** อย่างเด็ดขาด
- **ห้ามเขียน/สรุป/เรียบเรียงข้อความภาษาอังกฤษขึ้นมาใหม่** อย่างเด็ดขาด (ห้าม paraphrase)
- **ห้ามเอาข้อความภาษาอังกฤษจากข้ออื่นมาใส่** อย่างเด็ดขาด (เช่น ห้ามเอาของ S2 ไปใส่ให้ S3)
- **ต้องตรวจสอบเลขข้อให้ตรงกัน** ก่อนนำไปใช้
- ✅ **ข้อความอังกฤษที่ใส่ได้ ต้องเป็น "การคัดลอกตรงตัว (Exact Copy)" จากต้นฉบับเท่านั้น**
  - ตัวอักษร/คำ/เครื่องหมาย ต้องตรงกับที่ปรากฏ
  - ห้ามแต่งเพิ่ม ห้ามสรุป
- ถ้าหาข้อความภาษาอังกฤษไม่เจอ หรือ ไม่แน่ใจว่าตรงกัน -> **ต้องเว้นว่างไว้เท่านั้น** (ค่าว่างจริง ๆ)

## ⚠️ กฎเรื่อง (Rxx) - สำคัญมาก:
- ถ้าในต้นฉบับมี **(R1), (R2), (R3), (R101), (R111)** ฯลฯ นำหน้าข้อความอังกฤษ **ต้องใส่มาด้วยเสมอ**
- ตัวอย่าง: 
  - ต้นฉบับ: "(R101) MamyPoko Preemie/ Preterm / Small NB" 
  - Label_En: "(R101) MamyPoko Preemie/ Preterm / Small NB" ← ต้องใส่ (R101) มาด้วย
- **ห้ามตัด (Rxx) ออก** ถ้ามีในต้นฉบับ

## วิธีหาข้อความภาษาอังกฤษในแบบสอบถาม:
- แบบสอบถามเป็นตาราง 2 ภาษา โดยภาษาอังกฤษอาจอยู่:
  1) แถวถัดไป/ใต้ภาษาไทย
  2) ท้ายบรรทัดเดียวกัน
  3) คนละบรรทัดแต่ยังอยู่ใน "บล็อกของข้อเดียวกัน"
- **ต้องดูบล็อกที่มีเลขข้อตรงกับ `Variable` เท่านั้น**
- ตัวเลือกหนึ่ง ๆ ต้องจับคู่ ไทย (`Label_Th`) กับ อังกฤษที่อยู่ "ติดกัน/ใต้กัน" ภายในบล็อกเดียวกัน

## ขั้นตอนการทำงาน:

1) **ระบุคู่ข้อมูล:** สำหรับแต่ละแถว (`Variable`, `Value`, `Label_Th`):
   - หาบล็อกในแบบสอบถามที่มีเลขข้อตรงกับ `Variable` (เช่น s3 -> S3)
   - ในบล็อกนั้น ค้นหาตัวเลือกที่มีข้อความไทยตรงกับ `Label_Th`
   - จากตำแหน่งตัวเลือกนั้น ให้หา "ข้อความอังกฤษที่จับคู่กัน" (อยู่บรรทัดถัดไป/ท้ายบรรทัด/ใต้กัน ภายในตัวเลือกเดียวกัน)

2) **สร้างตารางผลลัพธ์:** สร้างตารางที่มี 3 คอลัมน์: `Variable`, `Value`, `Label_En`

3) **เติมข้อมูล (เข้มงวด):**
   - ใส่ `Variable` และ `Value` เหมือนเดิม
   - **เฉพาะกรณีที่พบข้อความอังกฤษที่เป็นคู่ของ `Label_Th` ภายในบล็อกเดียวกัน และคัดลอกได้ตรงตัวเท่านั้น:**
     - คัดลอกข้อความอังกฤษมาใส่ `Label_En` (Exact Copy)
     - **ถ้ามี (R1) (R2) (R3) (R101) ฯลฯ นำหน้าในต้นฉบับ ต้องใส่มาด้วยเสมอ**
   - **หากไม่เจอ หรือ ไม่แน่ใจว่าตรงกัน -> เว้นว่าง `Label_En`**
   - **ห้ามใส่ข้อความภาษาไทย** ลงใน `Label_En`

4) ✅ **Self-check ก่อนส่งคำตอบ (ห้ามข้าม):**
   - ตรวจทุกค่า `Label_En` ว่าสามารถพบเป็นข้อความเดียวกันในข้อมูลต้นฉบับ (Exact substring) จริงหรือไม่
   - ตรวจว่า (Rxx) ถูกใส่มาครบถ้ามีในต้นฉบับ
   - ถ้าไม่พบแบบตรงตัว -> ลบทิ้งแล้วเว้นว่าง

5) **จัดรูปแบบตาราง:**
   - ใช้ **Tab** เป็นตัวคั่นระหว่างคอลัมน์
   - ส่งกลับเป็น Table Markdown ที่มี 3 คอลัมน์: `Variable`, `Value`, `Label_En`
   - จัดให้อยู่ในรูปแบบ Codebox
"""

# --- Dark Mode Stylesheet (Larger Fonts) ---
DARK_STYLE = """
QMainWindow {
    background-color: #1e1e2e;
}
QWidget {
    background-color: #1e1e2e;
    color: #cdd6f4;
    font-family: 'Segoe UI', Tahoma;
    font-size: 14px;
}
QGroupBox {
    border: 1px solid #45475a;
    border-radius: 8px;
    margin-top: 16px;
    padding-top: 12px;
    font-weight: bold;
    font-size: 14px;
    color: #89b4fa;
}
QGroupBox::title {
    subcontrol-origin: margin;
    left: 10px;
    padding: 0 8px;
}
QPushButton {
    background-color: #45475a;
    color: #cdd6f4;
    border: 2px solid transparent;
    border-radius: 12px;
    padding: 14px 28px;
    font-weight: bold;
    font-size: 14px;
    min-height: 28px;
}
QPushButton:hover {
    background-color: #585b70;
    border: 2px solid #89b4fa;
}
QPushButton:pressed {
    background-color: #313244;
}
QPushButton:disabled {
    background-color: #313244;
    color: #6c7086;
}
QPushButton#btnGreen {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #a6e3a1, stop:1 #94e2d5);
    color: #1e1e2e;
    border: 2px solid #a6e3a1;
}
QPushButton#btnGreen:hover {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #94e2d5, stop:1 #89dceb);
    border: 2px solid #74c7ec;
}
QPushButton#btnBlue {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #89b4fa, stop:1 #74c7ec);
    color: #1e1e2e;
    border: 2px solid #89b4fa;
}
QPushButton#btnBlue:hover {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #74c7ec, stop:1 #89dceb);
    border: 2px solid #74c7ec;
}
QPushButton#btnPink {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #f38ba8, stop:1 #cba6f7);
    color: #1e1e2e;
    border: 2px solid #f38ba8;
}
QPushButton#btnPink:hover {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #eba0ac, stop:1 #ddb6fe);
    border: 2px solid #cba6f7;
}
QPushButton#btnOrange {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #fab387, stop:1 #f9e2af);
    color: #1e1e2e;
    border: 2px solid #fab387;
}
QPushButton#btnOrange:hover {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #f9e2af, stop:1 #f5c2e7);
    border: 2px solid #f9e2af;
}
QPushButton#btnTeal {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #94e2d5, stop:1 #89dceb);
    color: #1e1e2e;
    border: 2px solid #94e2d5;
}
QPushButton#btnTeal:hover {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #89dceb, stop:1 #74c7ec);
    border: 2px solid #89dceb;
}
QPushButton#btnPurple {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #cba6f7, stop:1 #b4befe);
    color: #1e1e2e;
    border: 2px solid #cba6f7;
}
QPushButton#btnPurple:hover {
    background-color: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #b4befe, stop:1 #89b4fa);
    border: 2px solid #b4befe;
}
QComboBox {
    background-color: #313244;
    border: 1px solid #45475a;
    border-radius: 6px;
    padding: 10px;
    font-size: 13px;
    min-width: 200px;
}
QComboBox::drop-down {
    border: none;
    width: 30px;
}
QComboBox::down-arrow {
    width: 12px;
    height: 12px;
}
QComboBox QAbstractItemView {
    background-color: #313244;
    selection-background-color: #45475a;
    font-size: 13px;
}
QTextEdit {
    background-color: #313244;
    border: 1px solid #45475a;
    border-radius: 8px;
    padding: 12px;
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 13px;
    line-height: 1.5;
}
QLabel {
    color: #cdd6f4;
    font-size: 13px;
}
QLabel#statusLabel {
    color: #a6e3a1;
    font-weight: bold;
    font-size: 15px;
    padding: 8px;
}
QLabel#fileLabel {
    color: #89b4fa;
    font-size: 12px;
}
QProgressBar {
    border: none;
    border-radius: 4px;
    background-color: #313244;
    height: 8px;
    text-align: center;
}
QProgressBar::chunk {
    background-color: #89b4fa;
    border-radius: 4px;
}
"""


class MissingValuesDialog(QDialog):
    """Dialog for user to manually fill in missing translations"""
    def __init__(self, missing_items, item_type="code", parent=None):
        super().__init__(parent)
        self.missing_items = missing_items  # List of dicts
        self.item_type = item_type  # "code" or "jod"
        self.user_inputs = {}
        self.setup_ui()
    
    def setup_ui(self):
        self.setWindowTitle(f"กรอกข้อมูลที่ขาด ({len(self.missing_items)} รายการ)")
        self.setMinimumSize(700, 500)
        self.setStyleSheet(DARK_STYLE)
        
        layout = QVBoxLayout(self)
        
        # Header
        header = QLabel(f"📝 พบ {len(self.missing_items)} รายการที่ไม่ได้แปล กรุณากรอกข้อมูลที่ขาด:")
        header.setStyleSheet("font-size: 14px; font-weight: bold; padding: 10px;")
        layout.addWidget(header)
        
        # Table
        self.table = QTableWidget()
        if self.item_type == "code":
            self.table.setColumnCount(4)
            self.table.setHorizontalHeaderLabels(["Variable", "Value", "Label_TH", "Label_EN (กรอกที่นี่)"])
        else:  # jod
            self.table.setColumnCount(3)
            self.table.setHorizontalHeaderLabels(["Name", "Label_TH", "VAR_ENG (กรอกที่นี่)"])
        
        self.table.setRowCount(len(self.missing_items))
        
        for row, item in enumerate(self.missing_items):
            if self.item_type == "code":
                self.table.setItem(row, 0, QTableWidgetItem(str(item.get('Variable', ''))))
                self.table.setItem(row, 1, QTableWidgetItem(str(item.get('Value', ''))))
                self.table.setItem(row, 2, QTableWidgetItem(str(item.get('Label_TH', ''))))
                self.table.setItem(row, 3, QTableWidgetItem(""))  # Editable
                # Make first 3 columns read-only
                for col in range(3):
                    self.table.item(row, col).setFlags(self.table.item(row, col).flags() & ~Qt.ItemFlag.ItemIsEditable)
            else:  # jod
                self.table.setItem(row, 0, QTableWidgetItem(str(item.get('Name', ''))))
                self.table.setItem(row, 1, QTableWidgetItem(str(item.get('Label_TH', ''))))
                self.table.setItem(row, 2, QTableWidgetItem(""))  # Editable
                # Make first 2 columns read-only
                for col in range(2):
                    self.table.item(row, col).setFlags(self.table.item(row, col).flags() & ~Qt.ItemFlag.ItemIsEditable)
        
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.resizeColumnsToContents()
        layout.addWidget(self.table)
        
        # Buttons
        btn_layout = QHBoxLayout()
        
        btn_ok = QPushButton("✅ ตกลง (บันทึก)")
        btn_ok.setStyleSheet("background-color: #a6e3a1; color: black; font-weight: bold; padding: 10px;")
        btn_ok.clicked.connect(self.accept)
        
        btn_cancel = QPushButton("❌ ยกเลิก (ปล่อยว่างไว้)")
        btn_cancel.setStyleSheet("background-color: #f38ba8; color: black; font-weight: bold; padding: 10px;")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
    
    def get_user_inputs(self):
        """Return dict of user-filled values"""
        results = {}
        for row in range(self.table.rowCount()):
            if self.item_type == "code":
                var = self.table.item(row, 0).text().strip().lower()
                val = self.table.item(row, 1).text().strip()
                label = self.table.item(row, 3).text().strip()  # User input column
                if label:
                    results[(var, val)] = label
            else:  # jod
                name = self.table.item(row, 0).text().strip().lower()
                var_eng = self.table.item(row, 2).text().strip()  # User input column
                if var_eng:
                    results[name] = var_eng
        return results


class ItemdefLoopDialog(QDialog):
    """Dialog for manual loop selection and Itemdef export with TB/T2B Making support"""
    
    # ค่าคงที่สำหรับ TB/T2B Making
    MAKING_OPTIONS_4 = ["T2B", "B2B"]  # 4 scale
    MAKING_OPTIONS_5 = ["TB", "T2B", "BB", "B2B"]  # 5 scale
    MAKING_OPTIONS_7_10 = ["TB", "T2B", "T3B", "BB", "B2B", "B3B"]  # 7/10 scale
    
    def __init__(self, spss_meta, spss_filepath, parent=None):
        super().__init__(parent)
        self.spss_meta = spss_meta
        self.spss_filepath = spss_filepath
        self.loop_groups = {}  # {var_name: group_name}
        self.scales_data = {}  # {var_name: {'num_options': n, 'direction': '', 'labels': []}}
        self.making_data = {}  # {var_name: {'conditions': [], 'making_var': 'Nvar'}}
        self.setup_ui()
        
    def setup_ui(self):
        self.setWindowTitle("กำหนด Loop สำหรับ Itemdef")
        self.setMinimumSize(900, 700)
        self.setStyleSheet(DARK_STYLE)
        
        layout = QVBoxLayout(self)
        
        # Instructions
        header = QLabel("1. เลือกตัวแปรที่เป็น Loop (ลากคลุมหลายบรรทัดได้)\n2. กดปุ่ม 'รวมกลุ่ม Loop' แล้วตั้งชื่อ (เช่น Q1)\n3. เมื่อเสร็จแล้วกด 'Export Itemdef' ด้านล่าง")
        header.setStyleSheet("font-size: 14px; font-weight: bold; color: #fab387; padding: 10px; background-color: #313244; border-radius: 8px;")
        layout.addWidget(header)
        
        # Tools Row 1 - Loop Grouping
        tool_layout = QHBoxLayout()
        
        btn_group = QPushButton("🔗 รวมกลุ่ม Loop (Group)")
        btn_group.setObjectName("btnBlue")
        btn_group.clicked.connect(self.group_selected)
        tool_layout.addWidget(btn_group)
        
        btn_ungroup = QPushButton("❌ ยกเลิกกลุ่ม (Ungroup)")
        btn_ungroup.setObjectName("btnPink")
        btn_ungroup.clicked.connect(self.ungroup_selected)
        tool_layout.addWidget(btn_ungroup)

        # NEW: Auto-Group MA
        btn_auto_ma = QPushButton("✨ Auto MA Groups")
        btn_auto_ma.setStyleSheet("background-color: #f38ba8; color: #1e1e2e; font-weight: bold;")
        btn_auto_ma.clicked.connect(self.auto_group_ma)
        tool_layout.addWidget(btn_auto_ma)
        
        tool_layout.addStretch()
        layout.addLayout(tool_layout)
        
        # Tools Row 2 - TB/T2B Making
        t2b_layout = QHBoxLayout()
        
        t2b_label = QLabel("TB/T2B Making:")
        t2b_label.setStyleSheet("font-weight: bold; color: #cba6f7;")
        t2b_layout.addWidget(t2b_label)
        
        btn_detect_scale = QPushButton("🔍 ตรวจจับ Scale 4/5/6/7/9/10")
        btn_detect_scale.setStyleSheet("background-color: #89b4fa; color: #1e1e2e; font-weight: bold;")
        btn_detect_scale.clicked.connect(self.detect_scale_variables)
        t2b_layout.addWidget(btn_detect_scale)
        
        btn_set_direction = QPushButton("➕ ตั้ง Direction (น้อยดี/มากดี)")
        btn_set_direction.setStyleSheet("background-color: #a6e3a1; color: #1e1e2e; font-weight: bold;")
        btn_set_direction.clicked.connect(self.set_direction_selected)
        t2b_layout.addWidget(btn_set_direction)
        
        btn_generate_making = QPushButton("⚡ สร้าง TB/T2B Making")
        btn_generate_making.setStyleSheet("background-color: #f9e2af; color: #1e1e2e; font-weight: bold;")
        btn_generate_making.clicked.connect(self.generate_making)
        t2b_layout.addWidget(btn_generate_making)
        
        t2b_layout.addStretch()
        layout.addLayout(t2b_layout)
        
        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(7)  # Added Direction column
        self.table.setHorizontalHeaderLabels([
            "No.", "Variable", "Label", "Type", 
            "Direction", "Loop Group", "Loop Type"
        ])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QTableWidget.SelectionMode.ContiguousSelection)
        self.table.cellDoubleClicked.connect(self.on_cell_double_click)
        
        layout.addWidget(self.table)
        
        # Load Data
        self.populate_table()
        
        # Auto-Run MA Grouping (Silent)
        self.auto_group_ma(silent=True)
        
        # Export Button
        btn_export = QPushButton("💾 Export Itemdef Excel (with Template)")
        btn_export.setObjectName("btnGreen")
        btn_export.setFixedHeight(50)
        btn_export.setFont(QFont("Segoe UI", 12, QFont.Weight.Bold))
        btn_export.clicked.connect(self.export_itemdef)
        layout.addWidget(btn_export)
        
        # Bottom Tools (Load/Save Config)
        config_layout = QHBoxLayout()
        btn_save_conf = QPushButton("บันทึก Config")
        btn_save_conf.clicked.connect(self.save_loop_config)
        config_layout.addWidget(btn_save_conf)
        
        btn_load_conf = QPushButton("โหลด Config")
        btn_load_conf.clicked.connect(self.load_loop_config)
        config_layout.addWidget(btn_load_conf)
        
        config_layout.addStretch()
        layout.addLayout(config_layout)
        
    def populate_table(self):
        self.table.setRowCount(0)
        if not self.spss_meta:
            return
            
        vars = self.spss_meta.column_names
        labels = self.spss_meta.column_labels
        readstat_types = getattr(self.spss_meta, 'readstat_variable_types', {}) 
        
        import re
        row_idx = 0
        
        for i, (var, label) in enumerate(zip(vars, labels)):
            # CHECK HIDDEN LOGIC (Only for detected MA groups)
            # If var matches _O\d+ (where d > 1) AND it belongs to a group of type MA (or Loop(MA) legacy) -> HIDE
            grp = self.loop_groups.get(var)
            grp_data = {}
            if isinstance(grp, dict): grp_data = grp
            elif isinstance(grp, str): grp_data = {'group': grp, 'type': 'MA'} # Assume detected legacy
            
            # Helper to check if it's _O2, _O3...
            is_ma_continuation = False
            match = re.search(r'_O(\d+)$', var)
            if match and int(match.group(1)) > 1:
                # Check if this variable is actually grouped as MA
                t = grp_data.get('type', '')
                if t == 'MA' or t == 'Loop(MA)':
                    is_ma_continuation = True
            
            if is_ma_continuation:
                continue # Skip adding row
            
            self.table.insertRow(row_idx)
            
            # --- Determine Type ---
            # If Grouped, use Group Type. If not, calculate Single Type.
            item_type = ""
            if grp_data.get('type'):
                 item_type = grp_data.get('type')
            else:
                 # Single Type Check
                 type_str = readstat_types.get(var, 'double')
                 item_type = "Numeric" 
                 if type_str == 'string': item_type = "Open-Ended"
                 else:
                     # Check labels -> SA
                      if self.spss_meta.variable_value_labels.get(var):
                           item_type = "SA"

            # No.
            self.table.setItem(row_idx, 0, QTableWidgetItem(str(i+1)))
            
            # Variable (Editable)
            var_item = QTableWidgetItem(var)
            var_item.setData(Qt.ItemDataRole.UserRole, var) # Store Original Var Name
            # Default is editable, but we want to ensure it is.
            # Only removing editable flag if it was restricted before.
            # To be safe: setFlags to Editable | Enabled | Selectable
            var_item.setFlags(Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEditable)
            self.table.setItem(row_idx, 1, var_item)
            
            # Label (Editable)
            lbl_item = QTableWidgetItem(str(label) if label else "")
            lbl_item.setFlags(Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable | Qt.ItemFlag.ItemIsEditable)
            self.table.setItem(row_idx, 2, lbl_item)
            
            # Type (Read-only)
            type_item = QTableWidgetItem(item_type)
            type_item.setFlags(type_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # Color Coding for Type
            if "MA" in item_type: type_item.setForeground(QColor("#a6e3a1"))
            elif "SA" in item_type: type_item.setForeground(QColor("#89b4fa"))
            elif "Open-Ended" in item_type: type_item.setForeground(QColor("#fab387"))
            
            self.table.setItem(row_idx, 3, type_item)
            
            # Direction (Col 4) - for TB/T2B Making
            direction_val = self.scales_data.get(var, {}).get('direction', '')
            scale_info = self.scales_data.get(var, {})
            
            # Build display text with first/last label hints
            display_text = direction_val
            if scale_info.get('num_options') in [4, 5, 6, 7, 9, 10]:
                labels_list = scale_info.get('labels', [])
                if labels_list and not direction_val:
                    first_lbl = labels_list[0][:15] if labels_list[0] else "?"
                    last_lbl = labels_list[-1][:15] if labels_list[-1] else "?"
                    display_text = f"[Scale {scale_info.get('num_options')}] 1={first_lbl}.. {scale_info.get('num_options')}={last_lbl}.."
                elif direction_val:
                    display_text = direction_val
            
            direction_item = QTableWidgetItem(display_text)
            direction_item.setFlags(direction_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # Color Direction based on scale detection
            if scale_info.get('num_options') in [4, 5, 6, 7, 9, 10]:
                if direction_val == "Scale น้อยดี(-)":
                    direction_item.setForeground(QColor("#f38ba8"))
                    direction_item.setBackground(QColor("#313244"))
                elif direction_val == "Scale มากดี(+)":
                    direction_item.setForeground(QColor("#a6e3a1"))
                    direction_item.setBackground(QColor("#313244"))
                else:
                    direction_item.setForeground(QColor("#cba6f7"))
                    direction_item.setBackground(QColor("#313244"))
            
            self.table.setItem(row_idx, 4, direction_item)
            
            # Group Info (Col 5, 6)
            grp_name = grp_data.get('group', "")
            grp_type = grp_data.get('type', "")

            group_item = QTableWidgetItem(grp_name)
            group_item.setFlags(group_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            loop_type_item = QTableWidgetItem(grp_type)
            loop_type_item.setFlags(loop_type_item.flags() & ~Qt.ItemFlag.ItemIsEditable)

            if grp_name:
                group_item.setBackground(QColor("#313244"))
                group_item.setForeground(QColor("#a6e3a1"))
                type_item.setBackground(QColor("#313244"))
                type_item.setForeground(QColor("#89b4fa"))

            self.table.setItem(row_idx, 5, group_item)
            self.table.setItem(row_idx, 6, loop_type_item)
            
            row_idx += 1

    def on_cell_double_click(self, row, column):
        """Double-click to view all value labels of a variable"""
        var_item = self.table.item(row, 1)
        if not var_item:
            return
        var = var_item.data(Qt.ItemDataRole.UserRole) or var_item.text()
        
        # Get value labels from SPSS metadata
        val_labels = self.spss_meta.variable_value_labels.get(var, {})
        if not val_labels:
            QMessageBox.information(self, f"Labels: {var}", 
                "ตัวแปรนี้ไม่มี Value Labels")
            return
        
        # Build display text
        codes = sorted([k for k in val_labels.keys() 
                       if isinstance(k, (int, float)) or 
                       (isinstance(k, str) and k.replace('.','',1).isdigit())],
                       key=lambda x: int(float(x)))
        
        text_lines = [f"Value Labels สำหรับ: {var}\n"]
        text_lines.append("-" * 40)
        for c in codes:
            lbl = val_labels.get(c) or val_labels.get(float(c)) or val_labels.get(str(c)) or ""
            text_lines.append(f"  {int(float(c)):>3} = {lbl}")
        text_lines.append("-" * 40)
        text_lines.append(f"\nจำนวน Labels: {len(codes)}")
        
        # Check if it's a scale variable
        if var in self.scales_data:
            direction = self.scales_data[var].get('direction', 'ยังไม่ตั้ง')
            text_lines.append(f"Direction: {direction}")
        
        QMessageBox.information(self, f"Labels: {var}", "\n".join(text_lines))

    def auto_group_ma(self, silent=False):
        """Auto detects _O1, _O2... variables and groups them as MA"""
        if not self.spss_meta: return
        
        import re
        vars = self.spss_meta.column_names
        ma_pattern = re.compile(r'(.+)_O(\d+)$')
        
        changes = 0
        processed_bases = set()
        
        for var in vars:
            match = ma_pattern.match(var)
            if match:
                base_name = match.group(1)
                # If we haven't processed this base yet
                if base_name not in processed_bases:
                    # Find all vars with this base
                    group_vars = [v for v in vars if re.match(rf'^{re.escape(base_name)}_O\d+$', v)]
                    if len(group_vars) > 1:
                        # Apply grouping
                        for v in group_vars:
                             self.loop_groups[v] = {'group': base_name, 'type': 'MA'}
                        changes += 1
                        processed_bases.add(base_name)
        
        if changes > 0:
            self.populate_table()
            if not silent: QMessageBox.information(self, "Auto Group", f"จับกลุ่ม MA ได้ {changes} กลุ่ม")
        else:
            if not silent: QMessageBox.information(self, "Auto Group", "ไม่พบกลุ่ม MA (Format: Name_O1, Name_O2...)")

    # ============ TB/T2B Making Functions ============
    
    def detect_scale_variables(self):
        """ตรวจจับตัวแปรที่มี Value Labels 4/5/6/7/9/10 options"""
        if not self.spss_meta:
            QMessageBox.warning(self, "เตือน", "กรุณาโหลดไฟล์ SPSS ก่อน")
            return
        
        vars_list = self.spss_meta.column_names
        found_scales = 0
        
        for var in vars_list:
            val_labels = self.spss_meta.variable_value_labels.get(var, {})
            if val_labels:
                # Count valid integer keys
                codes = [k for k in val_labels.keys() 
                         if isinstance(k, (int, float)) or 
                         (isinstance(k, str) and k.replace('.','',1).isdigit())]
                num_options = len(codes)
                
                if num_options in [4, 5, 6, 7, 9, 10]:
                    # Get labels in order
                    sorted_codes = sorted([int(float(c)) for c in codes])
                    labels_list = []
                    for c in sorted_codes:
                        lbl = val_labels.get(c) or val_labels.get(float(c)) or val_labels.get(str(c)) or ""
                        labels_list.append(str(lbl))
                    
                    self.scales_data[var] = {
                        'num_options': num_options,
                        'direction': self.scales_data.get(var, {}).get('direction', ''),
                        'labels': labels_list
                    }
                    found_scales += 1
        
        self.populate_table()
        QMessageBox.information(self, "ตรวจจับ Scale", 
            f"พบตัวแปร Scale 4/5/6/7/9/10 จำนวน {found_scales} ตัว\n\n"
            "คลิก 'ตั้ง Direction' เพื่อกำหนดทิศทาง Scale")
    
    def set_direction_selected(self):
        """ตั้งค่า Direction สำหรับตัวแปรที่เลือก"""
        selected_rows = sorted(set(idx.row() for idx in self.table.selectedIndexes()))
        if not selected_rows:
            QMessageBox.warning(self, "เตือน", "กรุณาเลือกตัวแปรที่จะตั้ง Direction ก่อน")
            return
        
        # Check if any selected is a valid scale
        valid_vars = []
        for row in selected_rows:
            var_item = self.table.item(row, 1)
            var = var_item.data(Qt.ItemDataRole.UserRole) or var_item.text()
            if var in self.scales_data:
                valid_vars.append(var)
        
        if not valid_vars:
            QMessageBox.warning(self, "เตือน", 
                "ตัวแปรที่เลือกไม่มี Scale 4/5/6/7/9/10\n\n"
                "กรุณากด 'ตรวจจับ Scale' ก่อน หรือเลือกตัวแปรที่มี Scale")
            return
        
        # Ask direction with custom colored dialog
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QPushButton
        
        dlg = QDialog(self)
        dlg.setWindowTitle("เลือก Direction")
        dlg.setMinimumWidth(280)
        dlg.setStyleSheet(DARK_STYLE)
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel(f"ตั้ง Direction สำหรับ {len(valid_vars)} ตัวแปร:"))
        
        selected_direction = [None]
        
        btn_style = "QPushButton {{ padding: 10px; font-size: 14px; font-weight: bold; border-radius: 5px; color: white; background-color: {bg}; }} QPushButton:hover {{ background-color: {hover}; }}"
        
        btn1 = QPushButton("🔴 Scale น้อยดี(-)")
        btn1.setStyleSheet(btn_style.format(bg="#e74c3c", hover="#c0392b"))
        btn1.clicked.connect(lambda: (selected_direction.__setitem__(0, "Scale น้อยดี(-)"), dlg.accept()))
        layout.addWidget(btn1)
        
        btn2 = QPushButton("🟢 Scale มากดี(+)")
        btn2.setStyleSheet(btn_style.format(bg="#27ae60", hover="#1e8449"))
        btn2.clicked.connect(lambda: (selected_direction.__setitem__(0, "Scale มากดี(+)"), dlg.accept()))
        layout.addWidget(btn2)
        
        btn3 = QPushButton("🟡 Justright(w+-) ไม่ทำTB")
        btn3.setStyleSheet(btn_style.format(bg="#f39c12", hover="#d68910"))
        btn3.clicked.connect(lambda: (selected_direction.__setitem__(0, "Justright(w+-) ไม่ทำTB"), dlg.accept()))
        layout.addWidget(btn3)
        
        btn4 = QPushButton("🟠 Justright(w+-) พร้อม TB")
        btn4.setStyleSheet(btn_style.format(bg="#e67e22", hover="#ca6f1e"))
        btn4.clicked.connect(lambda: (selected_direction.__setitem__(0, "Justright(w+-) พร้อม TB"), dlg.accept()))
        layout.addWidget(btn4)
        
        btn5 = QPushButton("🟣 Justright(ตามQNR) ไม่ทำTB")
        btn5.setStyleSheet(btn_style.format(bg="#8e44ad", hover="#71368a"))
        btn5.clicked.connect(lambda: (selected_direction.__setitem__(0, "Justright(ตามQNR) ไม่ทำTB"), dlg.accept()))
        layout.addWidget(btn5)
        
        btn6 = QPushButton("🟤 Justright(ตามQNR) พร้อม TB")
        btn6.setStyleSheet(btn_style.format(bg="#6d4c41", hover="#5d4037"))
        btn6.clicked.connect(lambda: (selected_direction.__setitem__(0, "Justright(ตามQNR) พร้อม TB"), dlg.accept()))
        layout.addWidget(btn6)
        
        btn7 = QPushButton("⬜ ยกเลิก Direction")
        btn7.setStyleSheet(btn_style.format(bg="#7f8c8d", hover="#5d6d7e"))
        btn7.clicked.connect(lambda: (selected_direction.__setitem__(0, ""), dlg.accept()))
        layout.addWidget(btn7)
        
        if dlg.exec() == QDialog.DialogCode.Accepted:
            direction = selected_direction[0]
            if direction is not None:
                for var in valid_vars:
                    self.scales_data[var]['direction'] = direction
                
                self.populate_table()
                if direction:
                    QMessageBox.information(self, "สำเร็จ", 
                        f"ตั้ง Direction '{direction}' ให้ {len(valid_vars)} ตัวแปร")
    
    def _generate_conditions(self, var, direction, num_labels):
        """สร้าง Conditions สำหรับ TB/T2B Making"""
        conditions = []
        justright_with_tb = {
            "Justright",
            "Justright(w+-) พร้อม TB",
            "Justright(ตามQNR) พร้อม TB",
        }
        
        # Original scale conditions
        for i in range(1, num_labels + 1):
            conditions.append(f"{var}={i}")
        
        # Making conditions based on direction
        if num_labels == 4:
            if direction == "Scale น้อยดี(-)":
                # T2B=1|2, B2B=3|4
                conditions.extend([
                    f"{var}=1|{var}=2",
                    f"{var}=3|{var}=4"
                ])
            elif direction == "Scale มากดี(+)":
                # T2B=4|3, B2B=1|2
                conditions.extend([
                    f"{var}=4|{var}=3",
                    f"{var}=1|{var}=2"
                ])
            elif direction in justright_with_tb:
                conditions.extend([
                    f"{var}=1|{var}=2",
                    f"{var}=3|{var}=4"
                ])
            else:
                conditions.extend(["NO_DIR"] * 2)
                
        elif num_labels in [5, 6]:
            if direction == "Scale น้อยดี(-)":
                # TB=1, T2B=1|2, BB=5, B2B=4|5
                high_code = num_labels
                conditions.extend([
                    f"{var}=1",
                    f"{var}=1|{var}=2",
                    f"{var}={high_code}",
                    f"{var}={high_code-1}|{var}={high_code}"
                ])
            elif direction == "Scale มากดี(+)":
                # TB=5, T2B=5|4, BB=1, B2B=1|2
                high_code = num_labels
                conditions.extend([
                    f"{var}={high_code}",
                    f"{var}={high_code}|{var}={high_code-1}",
                    f"{var}=1",
                    f"{var}=1|{var}=2"
                ])
            elif direction in justright_with_tb:
                # Only T2B(1+2) and B2B(4+5) - Scale มากดีเสมอ
                conditions.extend([
                    f"{var}=1|{var}=2",
                    f"{var}={num_labels-1}|{var}={num_labels}"
                ])
            else:
                conditions.extend(["NO_DIR"] * 4)
                
        elif num_labels in [7, 9, 10]:
            if direction == "Scale น้อยดี(-)":
                if num_labels == 7:
                    conditions.extend([
                        f"{var}=1",
                        f"{var}=1|{var}=2",
                        f"{var}=1|{var}=2|{var}=3",
                        f"{var}=7",
                        f"{var}=6|{var}=7",
                        f"{var}=5|{var}=6|{var}=7"
                    ])
                else:  # 9/10
                    high_code = num_labels
                    conditions.extend([
                        f"{var}=1",
                        f"{var}=1|{var}=2",
                        f"{var}=1|{var}=2|{var}=3",
                        f"{var}={high_code}",
                        f"{var}={high_code-1}|{var}={high_code}",
                        f"{var}={high_code-2}|{var}={high_code-1}|{var}={high_code}"
                    ])
            elif direction == "Scale มากดี(+)":
                if num_labels == 7:
                    conditions.extend([
                        f"{var}=7",
                        f"{var}=7|{var}=6",
                        f"{var}=7|{var}=6|{var}=5",
                        f"{var}=1",
                        f"{var}=1|{var}=2",
                        f"{var}=1|{var}=2|{var}=3"
                    ])
                else:  # 9/10
                    high_code = num_labels
                    conditions.extend([
                        f"{var}={high_code}",
                        f"{var}={high_code}|{var}={high_code-1}",
                        f"{var}={high_code}|{var}={high_code-1}|{var}={high_code-2}",
                        f"{var}=1",
                        f"{var}=1|{var}=2",
                        f"{var}=1|{var}=2|{var}=3"
                    ])
            elif direction in justright_with_tb:
                # Only T2B(1+2) and B2B - Scale มากดีเสมอ
                if num_labels == 7:
                    conditions.extend([
                        f"{var}=1|{var}=2",
                        f"{var}=6|{var}=7"
                    ])
                else:  # 9/10
                    high_code = num_labels
                    conditions.extend([
                        f"{var}=1|{var}=2",
                        f"{var}={high_code-1}|{var}={high_code}"
                    ])
            else:
                conditions.extend(["NO_DIR"] * 6)
        
        return conditions
    
    def generate_making(self):
        """สร้าง TB/T2B Making data"""
        if not self.scales_data:
            QMessageBox.warning(self, "เตือน", 
                "ยังไม่มีข้อมูล Scale\n\nกรุณากด 'ตรวจจับ Scale' ก่อน")
            return
        
        # Check variables with direction set
        valid_directions = [
            "Scale น้อยดี(-)",
            "Scale มากดี(+)",
            "Justright",
            "Justright(w+-) ไม่ทำTB",
            "Justright(w+-) พร้อม TB",
            "Justright(ตามQNR) ไม่ทำTB",
            "Justright(ตามQNR) พร้อม TB",
        ]
        vars_with_direction = [
            var for var, data in self.scales_data.items() 
            if data.get('direction') in valid_directions
        ]
        
        if not vars_with_direction:
            QMessageBox.warning(self, "เตือน", 
                "ยังไม่มีตัวแปรที่กำหนด Direction\n\n"
                "กรุณาเลือกตัวแปรและกด 'ตั้ง Direction' ก่อน")
            return
        
        # Generate making data
        self.making_data = {}
        skipped_no_tb = 0
        no_tb_directions = {"Justright(w+-) ไม่ทำTB", "Justright(ตามQNR) ไม่ทำTB"}
        for var in vars_with_direction:
            scale_info = self.scales_data[var]
            num_options = scale_info['num_options']
            direction = scale_info['direction']
            labels = scale_info.get('labels', [])
            
            if direction in no_tb_directions:
                skipped_no_tb += 1
                continue
            
            # Generate conditions
            conditions = self._generate_conditions(var, direction, num_options)
            
            # Making options based on scale and direction
            if direction in ["Justright", "Justright(w+-) พร้อม TB", "Justright(ตามQNR) พร้อม TB"]:
                # Only T2B(1+2) and B2B(4+5)
                if num_options == 7:
                    making_opts = ["(1+2)", "(6+7)"]
                elif num_options in [9, 10]:
                    making_opts = ["(1+2)", "(8+9)" if num_options == 9 else "(9+10)"]
                else:
                    making_opts = ["(1+2)", f"({num_options-1}+{num_options})"]
            elif num_options == 4:
                making_opts = self.MAKING_OPTIONS_4
            elif num_options in [5, 6]:
                making_opts = self.MAKING_OPTIONS_5
            else:
                making_opts = self.MAKING_OPTIONS_7_10
            
            self.making_data[var] = {
                'making_var': f"N{var}",
                'original_var': var,
                'num_options': num_options,
                'direction': direction,
                'conditions': conditions,
                'original_labels': labels,
                'making_labels': making_opts
            }
        
        msg = f"สร้าง TB/T2B Making สำหรับ {len(self.making_data)} ตัวแปร"
        if skipped_no_tb:
            msg += f"\n(ข้าม {skipped_no_tb} ตัวแปร: เลือกแบบไม่ทำTB)"
        QMessageBox.information(self, "สร้าง Making สำเร็จ", 
            f"{msg}\n\nข้อมูล Making จะถูกเพิ่มเมื่อ Export Itemdef")

    def group_selected(self):
        selected_rows = sorted(set(index.row() for index in self.table.selectedIndexes()))
        if not selected_rows:
            QMessageBox.warning(self, "เตือน", "กรุณาเลือกตัวแปรที่จะรวมกลุ่มก่อน")
            return
            
        # Get Group Name
        first_row = selected_rows[0]
        first_var = self.table.item(first_row, 1).text()
        
        # Suggest name
        suggest_name = first_var.rsplit('_', 1)[0]
        if not suggest_name: suggest_name = first_var
        
        from PyQt6.QtWidgets import QInputDialog
        group_name, ok = QInputDialog.getText(self, "ตั้งชื่อ Loop", "ชื่อ Loop Group (เช่น Q1):", text=suggest_name)
        
        if ok and group_name.strip():
            group_name = group_name.strip()
            
            # Ask for Loop Type
            items = ["Loop(SA)", "Loop(Text)", "Loop(Numeric)", "MA", "SA"]
            item, ok_type = QInputDialog.getItem(self, "เลือกประเภท Loop/Group", "Type:", items, 0, False)
            
            if ok_type and item:
                loop_type = item
                
                for row in selected_rows:
                    var = self.table.item(row, 1).data(Qt.ItemDataRole.UserRole) # Use original key
                    if not var: var = self.table.item(row, 1).text()
                    self.loop_groups[var] = {'group': group_name, 'type': loop_type}
                    
                self.populate_table()

    def ungroup_selected(self):
        selected_rows = sorted(set(index.row() for index in self.table.selectedIndexes()))
        if not selected_rows:
            return
            
        for row in selected_rows:
            var = self.table.item(row, 1).text()
            if var in self.loop_groups:
                del self.loop_groups[var]
        self.populate_table()

    def save_loop_config(self):
        """Save current loop grouping to Excel including Type"""
        if not self.loop_groups:
            QMessageBox.information(self, "ข้อมูล", "ยังไม่มีการกำหนดกลุ่ม Loop")
            return
            
        save_path, _ = QFileDialog.getSaveFileName(self, "บันทึก Config Loop", f"{os.path.splitext(os.path.basename(self.spss_filepath))[0]}_LoopConfig.xlsx", "Excel Files (*.xlsx)")
        if save_path:
            try:
                data = []
                for k, v in self.loop_groups.items():
                    if isinstance(v, dict):
                        data.append({"Variable": k, "Loop Group": v.get('group'), "Loop Type": v.get('type')})
                    else:
                        data.append({"Variable": k, "Loop Group": v, "Loop Type": "MA"})
                        
                df = pd.DataFrame(data)
                df.to_excel(save_path, index=False)
                QMessageBox.information(self, "สำเร็จ", "บันทึก Config เรียบร้อย")
            except Exception as e: QMessageBox.critical(self, "Error", f"{e}")

    def load_loop_config(self):
        """Load loop grouping from Excel"""
        load_path, _ = QFileDialog.getOpenFileName(self, "โหลด Config Loop", "", "Excel Files (*.xlsx)")
        if load_path:
            try:
                df = pd.read_excel(load_path)
                self.loop_groups = {}
                for _, row in df.iterrows():
                    var = str(row['Variable']).strip()
                    grp = str(row['Loop Group']).strip()
                    l_type = str(row.get('Loop Type', 'MA')).strip()
                    if var and grp and grp != "nan":
                         self.loop_groups[var] = {'group': grp, 'type': l_type}
                self.populate_table()
                QMessageBox.information(self, "สำเร็จ", f"โหลด Config เรียบร้อย ({len(self.loop_groups)} รายการ)")
            except Exception as e: QMessageBox.critical(self, "Error", f"{e}")

    def export_itemdef(self):
        base_dir = os.path.dirname(_resolve_resource_path("template.xlsx"))
        if not os.path.isdir(base_dir):
            base_dir = os.getcwd()
        # 1. Select Template
        template_path, _ = QFileDialog.getOpenFileName(
            self,
            "เลือกไฟล์ Template Format Itemdef",
            base_dir,
            "Excel Files (*.xlsx)"
        )
        if not template_path: return
        
        # 2. Select Output
        default_name = f"{os.path.splitext(os.path.basename(self.spss_filepath))[0]}_Itemdef.xlsx"
        default_path = os.path.join(base_dir, default_name)
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "บันทึก Itemdef Output",
            default_path,
            "Excel Files (*.xlsx)"
        )
        if not save_path: return

        # Show Loading Progress
        from PyQt6.QtWidgets import QProgressDialog
        from PyQt6.QtCore import Qt as QtCore_Qt
        progress = QProgressDialog("กำลัง Export Itemdef...", None, 0, 0, self)
        progress.setWindowTitle("กรุณารอสักครู่")
        progress.setWindowModality(QtCore_Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        try:
            # Build Table Overrides Map (for edited Rename/Label)
            # Map: original_var_name -> {'name': new_name, 'label': new_label}
            overrides = {}
            for row in range(self.table.rowCount()):
                 var_item_data = self.table.item(row, 1).data(Qt.ItemDataRole.UserRole)
                 original_var = var_item_data if var_item_data else self.table.item(row, 1).text()
                 
                 new_name = self.table.item(row, 1).text()
                 new_label = self.table.item(row, 2).text()
                 
                 if original_var:
                      overrides[original_var] = {'name': new_name, 'label': new_label}

            # Copy Template
            import shutil
            shutil.copyfile(template_path, save_path)
            
            wb = openpyxl.load_workbook(save_path)
            ws = wb.active # Assumption: Active sheet is the target
            
            # --- Column Indices (Based on Program_Itemdef logic) ---
            # Item=A, Fmt=B, Code=C, Type=D, Disp=E, Sub=F, ID=G, Label=H, ValID=G(reuse), ValLbl=H(reuse), Valid=I
            # We use 1-based index
            col_idx = {'Item':1, 'Fmt':2, 'Code':3, 'Type':4, 'Disp':5, 'LoopSub':6, 'ID':7, 'Label':8, 'ValID':7, 'ValLbl':8, 'Valid':9, 'Digit':11, 'Min':12, 'Max':13, 'Dec':14, 'Cat':10}
            
            start_row = 3 # Data starts at row 3 (Headings at 2?)
            current_row = start_row
            
            vars = self.spss_meta.column_names
            labels = self.spss_meta.column_labels
            readstat_types = getattr(self.spss_meta, 'readstat_variable_types', {}) 
            
            processed_groups = set()
            
            for index, var in enumerate(vars):
                # Check Overrides
                ov = overrides.get(var, {})
                final_var_name = ov.get('name', var)
                final_label = ov.get('label', labels[index])
                
                # Check if part of a group (use ORIGINAL var name for lookup)
                grp_info = self.loop_groups.get(var, {})
                
                if grp_info:
                    grp_name = grp_info['group']
                    grp_type = grp_info['type'] # Loop(SA), Loop(Text), MA...
                    
                    if grp_name in processed_groups: continue
                    
                    # New Loop Group Found
                    # Find all members
                    members = [v for v in vars if self.loop_groups.get(v, {}).get('group') == grp_name]
                    
                    # Handle "MA" separately as a Single Item (not Loop)
                    if grp_type == "MA" or grp_type == "SA": 
                         # Treat as Single Item but with Group Name
                         ws.cell(row=current_row, column=col_idx['Item'], value="Item")
                         ws.cell(row=current_row, column=col_idx['Fmt'], value="Survey")
                         ws.cell(row=current_row, column=col_idx['Type'], value=grp_type)
                         ws.cell(row=current_row, column=col_idx['Disp'], value="O")
                         ws.cell(row=current_row, column=col_idx['ID'], value=grp_name) # Use Group Name
                         
                         # Label from first member? Or Group Name?
                         # Usually use the label of the first member but strip _O1 info?
                         # Check overrides for the Base Member
                         first_mem_ov = overrides.get(members[0], {})
                         mem_label = first_mem_ov.get('label', labels[vars.index(members[0])] if members[0] in vars else "")
                         ws.cell(row=current_row, column=col_idx['Label'], value=mem_label)
                         
                         current_row += 1
                         
                         # Write Labels (from first member)
                         val_labels = self.spss_meta.variable_value_labels.get(members[0])
                         if val_labels:
                              codes = sorted([int(k) for k in val_labels.keys() if isinstance(k, (int, float)) or (isinstance(k, str) and k.isdigit())])
                              if codes:
                                   max_code = int(codes[-1])
                                   code_map = {}
                                   for k, v in val_labels.items():
                                        try: code_map[int(float(k))] = v
                                        except: pass
                                        
                                   for c in range(1, max_code + 1):
                                        ws.cell(row=current_row, column=col_idx['ValID'], value=c)
                                        ws.cell(row=current_row, column=col_idx['ValLbl'], value=code_map.get(c, ""))
                                        ws.cell(row=current_row, column=col_idx['Valid'], value="Valid")
                                        current_row += 1
                         
                    else:
                        # Real Loop
                        # 1. Write Header
                        ws.cell(row=current_row, column=col_idx['Item'], value="Item")
                        ws.cell(row=current_row, column=col_idx['Fmt'], value="Survey")
                        ws.cell(row=current_row, column=col_idx['Type'], value=grp_type)
                        ws.cell(row=current_row, column=col_idx['Disp'], value="O")
                        ws.cell(row=current_row, column=col_idx['ID'], value=grp_name)
                        
                        if "Numeric" in grp_type:
                             ws.cell(row=current_row, column=col_idx['Digit'], value=11)
                             ws.cell(row=current_row, column=col_idx['Min'], value=-9999999999)
                             ws.cell(row=current_row, column=col_idx['Max'], value=9999999999)
                             ws.cell(row=current_row, column=col_idx['Dec'], value=0)
                        elif "Text" in grp_type:
                             ws.cell(row=current_row, column=col_idx['Digit'], value=4000)

                        current_row += 1
                        
                        # 2. Write Sub Items
                        for i, mem_var in enumerate(members, 1):
                             mem_ov = overrides.get(mem_var, {})
                             mem_label = mem_ov.get('label', labels[vars.index(mem_var)] if mem_var in vars else "")
                             sub_id = f"{grp_name}({i})"
                             
                             ws.cell(row=current_row, column=col_idx['LoopSub'], value="Loop sub")
                             ws.cell(row=current_row, column=col_idx['ID'], value=sub_id)
                             ws.cell(row=current_row, column=col_idx['Label'], value=mem_label)
                             ws.cell(row=current_row, column=col_idx['Disp'], value="O")
                             current_row += 1
                             
                        # 3. Write Labels (From first member usually)
                        if "SA" in grp_type or "MA" in grp_type:
                             val_labels = self.spss_meta.variable_value_labels.get(members[0])
                             if val_labels:
                                  # Get integer keys
                                  codes = sorted([int(k) for k in val_labels.keys() if isinstance(k, (int, float)) or (isinstance(k, str) and k.isdigit())])
                                  if codes:
                                       max_code = int(codes[-1])
                                       # Map
                                       code_map = {}
                                       for k, v in val_labels.items():
                                            try: code_map[int(float(k))] = v
                                            except: pass
                                            
                                       for c in range(1, max_code + 1):
                                            ws.cell(row=current_row, column=col_idx['ValID'], value=c)
                                            ws.cell(row=current_row, column=col_idx['ValLbl'], value=code_map.get(c, ""))
                                            ws.cell(row=current_row, column=col_idx['Valid'], value="Valid")
                                            current_row += 1
                                        
                    processed_groups.add(grp_name)
                    
                else:
                    # Single Variable
                    # Mapping Readstat type to Itemdef format
                    type_str = readstat_types.get(var, 'double')
                    itemdef_type = "Numeric" 
                    if type_str == 'string': itemdef_type = "Open-Ended"
                    else:
                        # Check labels -> SA
                         if self.spss_meta.variable_value_labels.get(var):
                              itemdef_type = "SA"
                    
                    # Write Row
                    ws.cell(row=current_row, column=col_idx['Item'], value="Item")
                    ws.cell(row=current_row, column=col_idx['Fmt'], value="Survey")
                    ws.cell(row=current_row, column=col_idx['Type'], value=itemdef_type)
                    ws.cell(row=current_row, column=col_idx['Disp'], value="O")
                    ws.cell(row=current_row, column=col_idx['ID'], value=final_var_name) # Override
                    ws.cell(row=current_row, column=col_idx['Label'], value=final_label) # Override
                    
                    if itemdef_type == "Numeric":
                         ws.cell(row=current_row, column=col_idx['Digit'], value=11)
                         ws.cell(row=current_row, column=col_idx['Min'], value=-9999999999)
                         ws.cell(row=current_row, column=col_idx['Max'], value=9999999999)
                         ws.cell(row=current_row, column=col_idx['Dec'], value=0)
                    elif itemdef_type == "Open-Ended":
                         ws.cell(row=current_row, column=col_idx['Digit'], value=4000)
                    
                    current_row += 1
                    
                    # Write Labels if SA
                    if itemdef_type == "SA":
                         val_labels = self.spss_meta.variable_value_labels.get(var)
                         if val_labels:
                              codes = sorted([int(k) for k in val_labels.keys() if isinstance(k, (int, float)) or (isinstance(k, str) and k.isdigit())])
                              if codes:
                                   max_code = int(codes[-1])
                                   code_map = {}
                                   for k, v in val_labels.items():
                                        try: code_map[int(float(k))] = v
                                        except: pass
                                   
                                   # Calculate weights based on Direction (if any)
                                   weights = []
                                   scale_info = self.scales_data.get(var, {})
                                   direction = scale_info.get('direction', '')
                                   num_opts = scale_info.get('num_options', 0)
                                   if direction == "Scale น้อยดี(-)":
                                       weights = list(range(num_opts, 0, -1))
                                   elif direction == "Scale มากดี(+)":
                                       weights = list(range(1, num_opts + 1))
                                   elif direction in ["Justright", "Justright(ตามQNR) ไม่ทำTB", "Justright(ตามQNR) พร้อม TB"]:
                                       weights = list(range(1, num_opts + 1))
                                   elif direction in ["Justright(w+-) ไม่ทำTB", "Justright(w+-) พร้อม TB"]:
                                       half = num_opts // 2
                                       if num_opts % 2 == 1:
                                           weights = list(range(-half, 0)) + [0] + list(range(1, half + 1))
                                       else:
                                           weights = list(range(-half, 0)) + list(range(1, half + 1))
                                   
                                   for c in range(1, max_code + 1):
                                        ws.cell(row=current_row, column=col_idx['ValID'], value=c)
                                        ws.cell(row=current_row, column=col_idx['ValLbl'], value=code_map.get(c, ""))
                                        ws.cell(row=current_row, column=col_idx['Valid'], value="Valid")
                                        # Weight column (J=10) for SA labels
                                        if weights and c <= len(weights):
                                            ws.cell(row=current_row, column=10, value=weights[c-1])
                                        current_row += 1
                    
                    # ============ Write TB/T2B Making if exists ============
                    if var in self.making_data:
                        from openpyxl.styles import PatternFill
                        blue_fill = PatternFill(start_color="89B4FA", end_color="89B4FA", fill_type="solid")
                        
                        making_info = self.making_data[var]
                        making_var = making_info['making_var']
                        num_options = making_info['num_options']
                        direction = making_info['direction']
                        conditions = making_info['conditions']
                        original_labels = making_info['original_labels']
                        making_labels = making_info['making_labels']
                        
                        making_start_row = current_row
                        
                        # Write Making Header
                        ws.cell(row=current_row, column=col_idx['Item'], value="Item")
                        ws.cell(row=current_row, column=col_idx['Fmt'], value="Making")
                        ws.cell(row=current_row, column=col_idx['Type'], value="MA")
                        ws.cell(row=current_row, column=col_idx['Disp'], value="O")
                        ws.cell(row=current_row, column=col_idx['ID'], value=making_var)
                        ws.cell(row=current_row, column=col_idx['Label'], value=final_label)
                        # Statistic column (O=15) = original var name (keep case)
                        ws.cell(row=current_row, column=15, value=var)
                        # BaseType column (R=18)
                        ws.cell(row=current_row, column=18, value="Follow the condition items")
                        current_row += 1
                        
                        # Write Original Scale Options with conditions (no Weight here)
                        for i, lbl in enumerate(original_labels):
                            ws.cell(row=current_row, column=col_idx['ValID'], value=i+1)
                            ws.cell(row=current_row, column=col_idx['ValLbl'], value=lbl)
                            ws.cell(row=current_row, column=col_idx['Valid'], value="Valid")
                            # Conditions column (P=16)
                            if i < len(conditions):
                                ws.cell(row=current_row, column=16, value=conditions[i])
                            current_row += 1
                        
                        # Write TB/T2B/BB/B2B Options
                        start_cond_idx = num_options
                        for j, making_lbl in enumerate(making_labels):
                            ws.cell(row=current_row, column=col_idx['ValID'], value=num_options + j + 1)
                            ws.cell(row=current_row, column=col_idx['ValLbl'], value=making_lbl)
                            ws.cell(row=current_row, column=col_idx['Valid'], value="Valid")
                            cond_idx = start_cond_idx + j
                            if cond_idx < len(conditions):
                                ws.cell(row=current_row, column=16, value=conditions[cond_idx])
                            current_row += 1
                        
                        # Apply blue fill to all Making rows
                        for r in range(making_start_row, current_row):
                            for c in range(1, 20):
                                ws.cell(row=r, column=c).fill = blue_fill

            # Write END marker
            ws.cell(row=current_row, column=col_idx['Item'], value="End")

            wb.save(save_path)
            progress.close()
            QMessageBox.information(self, "สำเร็จ", f"บันทึกไฟล์เรียบร้อยแล้ว:\n{save_path}")
            self.accept()
            
        except Exception as e:
            progress.close()
            QMessageBox.critical(self, "Error", f"เกิดข้อผิดพลาดในการ Export:\n{e}")
            import traceback
            traceback.print_exc()

class AIWorker(QThread):
    """Thread สำหรับเรียก API แบบ non-blocking"""
    progress = pyqtSignal(str)  # ส่ง content chunk
    finished = pyqtSignal(bool, str, int, int)  # success, message, tokens, time_ms
    
    def __init__(self, prompt, model):
        super().__init__()
        self.prompt = prompt
        self.model = model
        self.result_text = ""
        
    def run(self):
        start_time = time.time()
        total_tokens = 0
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                # Set up session with keep-alive
                session = requests.Session()
                
                api_key = get_openrouter_api_key()
                if not api_key:
                    elapsed_ms = int((time.time() - start_time) * 1000)
                    self.finished.emit(False, "Missing OpenRouter API key.", 0, elapsed_ms)
                    return

                response = session.post(
                    url="https://openrouter.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "https://github.com/GenPromt/App", # Best practice for OpenRouter
                        "X-Title": "GenPromt App"
                    },
                    data=json.dumps({
                        "model": self.model,
                        "messages": [{"role": "user", "content": self.prompt}],
                        "include_reasoning": True,
                        "stream": True # Always stream
                    }),
                    stream=True,
                    timeout=900 # 15 minutes timeout
                )
                
                if response.status_code == 200:
                    for line in response.iter_lines():
                        if line:
                            line_str = line.decode('utf-8')
                            if line_str.startswith('data: '):
                                data_str = line_str[6:]
                                if data_str == '[DONE]':
                                    break
                                try:
                                    data = json.loads(data_str)
                                    choice = data.get('choices', [{}])[0]
                                    delta = choice.get('delta', {})
                                    
                                    # Handle Reasoning (DeepSeek R1)
                                    reasoning_chunk = delta.get('reasoning', '')
                                    if reasoning_chunk:
                                        self.progress.emit(f"<span style='color: #6c7086;'>{reasoning_chunk}</span>")

                                    content_chunk = delta.get('content', '')
                                    if content_chunk:
                                        self.result_text += content_chunk
                                        self.progress.emit(content_chunk)
                                    
                                    # Get token usage
                                    usage = data.get('usage', {})
                                    if usage:
                                        total_tokens = usage.get('total_tokens', 0)
                                except json.JSONDecodeError:
                                    pass
                    
                    elapsed_ms = int((time.time() - start_time) * 1000)
                    self.finished.emit(True, self.result_text, total_tokens, elapsed_ms)
                    return # Success, exit loop
                    
                elif response.status_code == 429:
                    # Rate limit - wait and retry? Or just fail?
                    # Generally wait a bit
                    if attempt < max_retries - 1:
                        time.sleep(5) # Wait 5 sec before retry
                        continue
                    else:
                         elapsed_ms = int((time.time() - start_time) * 1000)
                         self.finished.emit(False, f"Rate Limited (429). Please try a different model.", 0, elapsed_ms)
                         return

                else:
                    # Other non-200 errors
                    if attempt < max_retries - 1:
                        time.sleep(2)
                        continue
                    elapsed_ms = int((time.time() - start_time) * 1000)
                    self.finished.emit(False, f"API Error ({response.status_code}): {response.text}", 0, elapsed_ms)
                    return
                    
            except (requests.exceptions.ConnectionError, 
                    requests.exceptions.Timeout, 
                    requests.exceptions.ChunkedEncodingError,
                    requests.exceptions.RequestException) as e:
                
                if attempt < max_retries - 1:
                    # Notify visible progress of retry
                    self.progress.emit(f"\n\n<i>⚠️ Connection lost. Retrying ({attempt+1}/{max_retries})...</i>\n\n")
                    time.sleep(3) # Wait before retry
                    continue
                else:
                    elapsed_ms = int((time.time() - start_time) * 1000)
                    self.finished.emit(False, f"Failed after {max_retries} retries. Error: {str(e)}", 0, elapsed_ms)
                    return
            except Exception as e:
                 elapsed_ms = int((time.time() - start_time) * 1000)
                 self.finished.emit(False, str(e), 0, elapsed_ms)
                 return


class GenPromtApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("FullGenpromt+Itemdef Beta V1")
        self.setMinimumSize(1000, 750)
        
        # Data storage
        self.questionnaire_data = ""
        self.cat_program_data = {"var": "", "value": ""}
        self.cat_program_filepath = ""
        self.current_prompt_type = None
        self.ai_result_text = ""
        self.is_all_mode = False
        self.jod_result_data = []
        self.total_tokens_used = 0
        self.session_total_tokens = 0
        self.session_total_time_ms = 0
        
        # AI Worker
        self.worker = None

        # Timer for realtime updates
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_realtime_stats)
        self.current_chunk_start_time = 0
        
        self.init_ui()
        QTimer.singleShot(0, self.center_window)

    def center_window(self):
        screen = QApplication.primaryScreen()
        if not screen:
            return
        frame = self.frameGeometry()
        frame.moveCenter(screen.availableGeometry().center())
        self.move(frame.topLeft())

    def update_realtime_stats(self):
        if not hasattr(self, 'current_chunk_start_time') or self.current_chunk_start_time == 0:
            return
            
        elapsed_chunk_ms = int((time.time() - self.current_chunk_start_time) * 1000)
        total_session_ms = self.session_total_time_ms + elapsed_chunk_ms
        
        # Format time
        minutes = total_session_ms // 60000
        seconds = (total_session_ms % 60000) // 1000
        if minutes > 0:
            time_str = f"{minutes} นาที {seconds} วินาที"
        else:
            time_str = f"{seconds} วินาที"
            
        self.lbl_stats.setText(f"⏱ {time_str} (Running) | 📊 {self.session_total_tokens:,} tokens")

    def set_buttons_enabled(self, enabled):
        """Enable or disable action buttons"""
        self.btn_prompt_jod.setEnabled(enabled)
        self.btn_prompt_code.setEnabled(enabled)
        self.btn_prompt_all.setEnabled(enabled)
        self.btn_download.setEnabled(enabled)
        if enabled:
            # Enable SPSS save button only after processing
            self.btn_save_spss.setEnabled(True)
    
    def start_ai(self, prompt, mode="normal"):
        self.set_buttons_enabled(False)
        if mode == "normal" or (mode == "code_chunk" and self.current_chunk_index == 0):
            self.output_text.clear()
            
        if mode == "normal":
            self.lbl_status.setText("⏳ กำลังประมวลผล...")
            
        self.lbl_status.setStyleSheet("color: #89b4fa; font-weight: bold;")
        self.current_ai_mode = mode # Store mode to handle callback
        
        # Start Timer
        self.current_chunk_start_time = time.time()
        self.timer.start(1000) # Update every 1 second
        
        model = self.model_map.get(self.model_combo.currentText(), self.model_combo.currentText())
        self.worker = AIWorker(prompt, model)
        self.worker.progress.connect(self.on_ai_progress)
        self.worker.finished.connect(self.on_ai_finished)
        self.worker.start()
        
    def on_ai_finished(self, success, result, tokens, time_ms):
        self.timer.stop() # Stop timer first
        self.current_chunk_start_time = 0
        
        self.ai_result_text = result # Will be overwritten if chunking
        
        # Accumulate stats
        if not hasattr(self, 'session_total_tokens'):
             self.session_total_tokens = 0
             self.session_total_time_ms = 0
        
        if success:
             self.session_total_tokens += tokens
             self.session_total_time_ms += time_ms
        
        self.total_tokens_used = self.session_total_tokens
        
        # Format TOTAL time
        minutes = self.session_total_time_ms // 60000
        seconds = (self.session_total_time_ms % 60000) // 1000
        if minutes > 0:
            time_str = f"{minutes} นาที {seconds} วินาที"
        else:
            time_str = f"{seconds} วินาที"
        
        if success:
            # Handle Code Chunking
            if hasattr(self, 'current_ai_mode') and self.current_ai_mode == "code_chunk":
                self.accumulated_code_results.append(result)
                self.current_chunk_index += 1
                
                # Check if more chunks
                val_lines = self.cat_program_data['value'].strip().split('\n')
                header = val_lines[0]
                
                if self.current_chunk_index < len(self.code_chunks):
                    # Continue to next chunk
                    self.process_next_code_chunk(header)
                    return
                else:
                    # All chunks done
                    self.process_next_code_chunk(header)
                    self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens")
                    return

            # Check ALL mode (JOD finished)
            if self.is_all_mode and self.current_prompt_type == "jod":
                # Save JOD result and continue with Code
                self.jod_result_data = self.parse_ai_result("jod")
                self.lbl_status.setText("✅ Prompt โจทย์ เสร็จ → กำลังรัน Code...")
                # run_prompt_code will start timer again
                
                # Run prompt code (will handle chunking internally)
                self.run_prompt_code() # This initiates chunking
                return
                
            elif self.is_all_mode and self.current_prompt_type == "code":
                 pass

            if not hasattr(self, 'current_ai_mode') or self.current_ai_mode == "normal":
                 if self.is_all_mode:
                      pass
                 else:
                      self.lbl_status.setText("✅ เสร็จสิ้น!")

            self.lbl_status.setStyleSheet("color: #a6e3a1; font-weight: bold;")
            self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens")
        else:
            self.is_all_mode = False
            self.lbl_status.setText("❌ เกิดข้อผิดพลาด")
            self.lbl_status.setStyleSheet("color: #f38ba8; font-weight: bold;")
            self.lbl_stats.setText(f"⏱ {time_str}")
            QMessageBox.critical(self, "Error", result)
        
    def init_ui(self):
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # === Header ===
        header = QLabel("🔷 Full Genpromt+Itemdef+T2B Beta V1")
        header.setStyleSheet("font-size: 18px; font-weight: bold; color: #89b4fa; padding: 10px;")
        header.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(header)
        
        # === File Selection Area ===
        file_layout = QHBoxLayout()
        
        # Questionnaire file group
        quest_group = QGroupBox("📄 ไฟล์แบบสอบถาม")
        quest_layout = QVBoxLayout()
        self.btn_select_quest = QPushButton("📂 เลือกไฟล์...")
        self.btn_select_quest.clicked.connect(self.select_questionnaire)
        quest_layout.addWidget(self.btn_select_quest)
        self.lbl_quest_file = QLabel("ยังไม่ได้เลือกไฟล์")
        self.lbl_quest_file.setObjectName("fileLabel")
        self.lbl_quest_file.setWordWrap(True)
        quest_layout.addWidget(self.lbl_quest_file)
        quest_group.setLayout(quest_layout)
        file_layout.addWidget(quest_group)
        
        # SPSS file group (เปลี่ยนจาก Excel เป็น SPSS)
        spss_group = QGroupBox("📊 ไฟล์ SPSS (.sav)")
        spss_layout = QVBoxLayout()
        
        # MA Delimiter selector
        ma_layout = QHBoxLayout()
        ma_label = QLabel("MA Delimiter:")
        ma_label.setStyleSheet("font-size: 12px;")
        ma_layout.addWidget(ma_label)
        self.ma_combo = QComboBox()
        self.ma_combo.addItems(["_O", "$"])
        self.ma_combo.setFixedWidth(60)
        ma_layout.addWidget(self.ma_combo)
        ma_layout.addStretch()
        spss_layout.addLayout(ma_layout)
        
        self.btn_select_spss = QPushButton("📂 เลือกไฟล์ SPSS...")
        self.btn_select_spss.clicked.connect(self.load_spss_file)
        spss_layout.addWidget(self.btn_select_spss)
        
        self.lbl_spss_file = QLabel("ยังไม่ได้เลือกไฟล์")
        self.lbl_spss_file.setObjectName("fileLabel")
        self.lbl_spss_file.setWordWrap(True)
        spss_layout.addWidget(self.lbl_spss_file)
        
        self.lbl_spss_status = QLabel("")
        self.lbl_spss_status.setStyleSheet("color: #a6e3a1; font-size: 12px;")
        spss_layout.addWidget(self.lbl_spss_status)
        
        spss_group.setLayout(spss_layout)
        file_layout.addWidget(spss_group)
        
        main_layout.addLayout(file_layout)
        
        # === Model Selection ===
        model_layout = QHBoxLayout()
        model_label = QLabel("🤖 Agent:")
        model_label.setStyleSheet("font-weight: bold;")
        model_layout.addWidget(model_label)
        
        self.model_combo = QComboBox()
        self.model_combo.addItems([
            "🧑‍💼 พนักงานตัวน้อย",
            "👔 CEO"
        ])
        self.model_combo.setCurrentIndex(1)  # ตั้ง "👔 CEO" เป็น default
        # Map display name to actual model
        self.model_map = {
            "🧑‍💼 พนักงานตัวน้อย": "google/gemini-2.5-flash-lite",
            "👔 CEO": "google/gemini-3-flash-preview"
        }
        model_layout.addWidget(self.model_combo)
        model_layout.addStretch()
        main_layout.addLayout(model_layout)
        
        # === Action Buttons ===
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(15)
        
        self.btn_prompt_jod = QPushButton("🟢 Prompt โจทย์")
        self.btn_prompt_jod.setObjectName("btnGreen")
        self.btn_prompt_jod.clicked.connect(self.run_prompt_jod)
        btn_layout.addWidget(self.btn_prompt_jod)
        
        self.btn_prompt_code = QPushButton("🔵 Prompt Code")
        self.btn_prompt_code.setObjectName("btnBlue")
        self.btn_prompt_code.clicked.connect(self.run_prompt_code)
        btn_layout.addWidget(self.btn_prompt_code)
        
        self.btn_prompt_all = QPushButton("🔴 Prompt ALL")
        self.btn_prompt_all.setObjectName("btnPink")
        self.btn_prompt_all.clicked.connect(self.run_prompt_all)
        btn_layout.addWidget(self.btn_prompt_all)
        
        main_layout.addLayout(btn_layout)
        
        # === Output Area ===
        output_group = QGroupBox("📝 ผลลัพธ์")
        output_layout = QVBoxLayout()
        
        self.output_text = QTextEdit()
        self.output_text.setAcceptRichText(True)  # Enable HTML for reasoning
        self.output_text.setPlaceholderText("ผลลัพธ์จะแสดงที่นี่...")
        self.output_text.setMinimumHeight(300)
        output_layout.addWidget(self.output_text)
        
        # Status bar
        status_layout = QHBoxLayout()
        self.lbl_status = QLabel("พร้อมใช้งาน")
        self.lbl_status.setObjectName("statusLabel")
        status_layout.addWidget(self.lbl_status)
        status_layout.addStretch()
        
        self.lbl_stats = QLabel("")
        self.lbl_stats.setStyleSheet("color: #f9e2af; font-size: 14px; font-weight: bold;")
        status_layout.addWidget(self.lbl_stats)
        output_layout.addLayout(status_layout)
        
        output_group.setLayout(output_layout)
        main_layout.addWidget(output_group)
        
        # === Bottom Buttons ===
        bottom_layout = QHBoxLayout()
        bottom_layout.setSpacing(15)
        
        self.btn_copy = QPushButton("📋 คัดลอกผลลัพธ์")
        self.btn_copy.clicked.connect(self.copy_to_clipboard)
        bottom_layout.addWidget(self.btn_copy)
        
        self.btn_download = QPushButton("💾 Download Excel")
        self.btn_download.setObjectName("btnTeal")
        self.btn_download.clicked.connect(self.download_excel)
        bottom_layout.addWidget(self.btn_download)
        
        self.btn_load_excel = QPushButton("📂 โหลดไฟล์ Excel")
        self.btn_load_excel.setObjectName("btnOrange")
        self.btn_load_excel.clicked.connect(self.load_ai_result_excel)
        bottom_layout.addWidget(self.btn_load_excel)
        
        self.btn_save_spss = QPushButton("💾 Save to SPSS")
        self.btn_save_spss.setObjectName("btnPurple")
        self.btn_save_spss.clicked.connect(self.save_labels_to_spss)
        self.btn_save_spss.setEnabled(False)
        bottom_layout.addWidget(self.btn_save_spss)
        
        self.btn_itemdef = QPushButton("สร้าง Itemdef 📁")
        self.btn_itemdef.setStyleSheet("background-color: #f9e2af; color: #1e1e2e; font-weight: bold; padding: 12px 24px; border-radius: 8px;")
        self.btn_itemdef.clicked.connect(self.open_itemdef_dialog)
        self.btn_itemdef.setEnabled(False) # Enabled only when SPSS loaded
        bottom_layout.addWidget(self.btn_itemdef)
        
        bottom_layout.addStretch()
        main_layout.addLayout(bottom_layout)
        
    # === File Selection Methods ===
    def select_questionnaire(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "เลือกไฟล์แบบสอบถาม", "",
            "Word/Excel Files (*.docx *.xlsx);;All Files (*.*)"
        )
        if files:
            all_text = []
            for filepath in files:
                try:
                    if filepath.endswith('.docx'):
                        doc = Document(filepath)
                        
                        # === ดึงข้อมูลจาก PARAGRAPHS ===
                        for p in doc.paragraphs:
                            if p.text.strip():
                                all_text.append(p.text)
                        
                        # === ดึงข้อมูลจาก TABLES (สำคัญมาก!) ===
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    # ดึงข้อความทั้งหมดในเซลล์ (รวมหลายบรรทัด)
                                    cell_text = cell.text.strip()
                                    if cell_text:
                                        row_text.append(cell_text)
                                if row_text:
                                    all_text.append('\t'.join(row_text))
                                    
                    elif filepath.endswith('.xlsx'):
                        wb = openpyxl.load_workbook(filepath, data_only=True)
                        for sheet in wb.worksheets:
                            for row in sheet.iter_rows():
                                row_text = [str(cell.value or '') for cell in row]
                                all_text.append('\t'.join(row_text))
                except Exception as e:
                    QMessageBox.warning(self, "Error", f"ไม่สามารถอ่านไฟล์ได้: {e}")
                    return
            
            self.questionnaire_data = '\n'.join(all_text)
            names = [os.path.basename(f) for f in files]
            self.lbl_quest_file.setText('\n'.join(names))
            
            # === แสดงจำนวนข้อมูลที่โหลดได้ ===
            line_count = len(all_text)
            char_count = len(self.questionnaire_data)
            print(f"✅ Loaded questionnaire: {line_count} lines, {char_count:,} characters")
            
    def load_spss_file(self):
        spss_path, _ = QFileDialog.getOpenFileName(
            self, "เลือกไฟล์ SPSS", "",
            "SPSS Files (*.sav);;All Files (*.*)"
        )
        if spss_path:
            self.lbl_spss_status.setText("⏳ กำลังอ่านไฟล์ SPSS...")
            QApplication.processEvents()
            
            try:
                delimiter = self.ma_combo.currentText()
                df, meta, df_var, df_value = self._process_spss_data(spss_path, delimiter)
                
                if df is not None:
                    # Store for SPSS save feature
                    self.spss_filepath = spss_path
                    self.spss_meta = meta
                    self.original_spss_df = df
                    
                    # Auto save to temp excel
                    temp_dir = tempfile.gettempdir()
                    self.cat_program_filepath = os.path.join(temp_dir, "temp_cat_program.xlsx")
                    
                    with pd.ExcelWriter(self.cat_program_filepath, engine="openpyxl") as writer:
                        df_var.to_excel(writer, index=False, sheet_name='Var')
                        df_value.to_excel(writer, index=False, sheet_name='Value')
                    
                    # Compute tab-separated text for prompts
                    # Var sheet (send Name + VAR_THA)
                    var_lines = []
                    var_lines.append(f"{df_var.columns[0]}\t{df_var.columns[1]}")
                    for _, row in df_var.iterrows():
                        var_lines.append(f"{row['Name']}\t{row['VAR_THA']}")
                    self.cat_program_data['var'] = '\n'.join(var_lines)
                    
                    # Value sheet (send Variable + Value + Label_Th)
                    val_lines = []
                    val_lines.append(f"{df_value.columns[0]}\t{df_value.columns[1]}\t{df_value.columns[2]}")
                    for _, row in df_value.iterrows():
                        val_lines.append(f"{row['Variable']}\t{row['Value']}\t{row['Label_Th']}")
                    self.cat_program_data['value'] = '\n'.join(val_lines)
                    
                    self.lbl_spss_file.setText(os.path.basename(spss_path))
                    self.lbl_spss_status.setText(f"✅ โหลดสำเร็จ (Var: {len(df_var)}, Value: {len(df_value)})")

                    # Enable buttons dependent on SPSS
                    if hasattr(self, 'btn_itemdef'):
                        self.btn_itemdef.setEnabled(True)
                
            except Exception as e:
                self.lbl_spss_status.setText("❌ เกิดข้อผิดพลาด")
                QMessageBox.critical(self, "Error", f"อ่านไฟล์ SPSS ไม่สำเร็จ: {e}")

    def open_itemdef_dialog(self):
        if not hasattr(self, 'spss_meta') or not self.spss_meta:
             QMessageBox.warning(self, "เตือน", "กรุณาโหลดไฟล์ SPSS ก่อน")
             return
             
        dialog = ItemdefLoopDialog(self.spss_meta, self.spss_filepath, self)
        dialog.exec()

    def _process_spss_data(self, spss_file_path, delimiter):
        try:
            df, meta = pyreadstat.read_sav(spss_file_path)
            
            label_data = []
            for var_name, var_labels in meta.variable_value_labels.items():
                if delimiter in var_name and not var_name.endswith(f"{delimiter}1"):
                    continue
                for value, label in var_labels.items(): 
                    label_data.append({'Variable': var_name, 'Value': value, 'Label_Th': label})
            
            labels_df = pd.DataFrame(label_data)
            if not labels_df.empty:
                labels_df['Label_EN'] = ""

            filtered_column_names = [name for name in meta.column_names if delimiter not in name or name.endswith(f"{delimiter}1")]
            filtered_column_labels = [label for name, label in zip(meta.column_names, meta.column_labels) if delimiter not in name or name.endswith(f"{delimiter}1")]

            labels_dd = pd.DataFrame({
                'Name': filtered_column_names,
                'VAR_THA': filtered_column_labels
            })
            labels_dd['VAR_ENG'] = ""
            
            return df, meta, labels_dd, labels_df
        except Exception as e:
            raise e

    # === Prompt Methods ===
    def run_prompt_jod(self):
        if not self.validate_inputs("jod"):
            return
        
        # Reset Session Stats
        self.session_total_tokens = 0
        self.session_total_time_ms = 0
        
        # Initialize two-tier processing for Jod
        self.jod_two_tier_pass = 1
        self.jod_first_pass_results = {}  # {name: var_eng}
        self.jod_empty_rows_for_second_pass = []
        
        # Prepare data
        var_lines = self.cat_program_data['var'].strip().split('\n')
        self.jod_header = var_lines[0]
        self.jod_data_rows = var_lines[1:]
        
        # Start Pass 1 with cheap model
        self.run_jod_pass(1)
    
    def run_jod_pass(self, pass_number):
        """Run a single pass of Jod (Variable) translation"""
        self.jod_two_tier_pass = pass_number
        
        selected_model = self.model_map.get(self.model_combo.currentText(), self.model_combo.currentText())
        is_lite_mode = (selected_model == "google/gemini-2.5-flash-lite")
        
        if pass_number == 1:
            # Pass 1
            if is_lite_mode:
                model = "google/gemini-2.5-flash-lite"
                status_msg = "⏳ โจทย์ Pass 1/2: พนักงานตัวน้อยกำลังทำงาน..."
            else:
                model = selected_model
                # แสดงชื่อ Agent จาก dropdown แทน model name
                agent_name = self.model_combo.currentText()
                status_msg = f"⏳ โจทย์ Pass 1/2: ใช้ {agent_name}ทำงาน..."
                
            data_text = self.jod_header + '\n' + '\n'.join(self.jod_data_rows)
        else:
            # Pass 2
            if not self.jod_empty_rows_for_second_pass:
                self.finish_jod_two_tier_processing()
                return
                
            if is_lite_mode:
                model = "google/gemini-3-flash-preview" # Upgrade for pass 2 if in Lite mode
                status_msg = f"⏳ โจทย์ Pass 2/2: CEO กำลังดูงาน สำหรับ {len(self.jod_empty_rows_for_second_pass)} แถวที่ว่าง..."
            else:
                model = selected_model # Stick to selected model
                agent_name = self.model_combo.currentText()
                status_msg = f"⏳ โจทย์ Pass 2/2: ใช้ {agent_name} เก็บตก {len(self.jod_empty_rows_for_second_pass)} แถว..."
                
            data_text = self.jod_header + '\n' + '\n'.join(self.jod_empty_rows_for_second_pass)
        
        combined_data = f"{self.questionnaire_data}\n\n{data_text}"
        prompt = PROMPT_JOD.replace("[ *** วางข้อมูล Word และ Excel *** ]", combined_data)
        
        self.current_prompt_type = "jod"
        self.lbl_status.setText(status_msg)
        self.lbl_status.setStyleSheet("color: #89b4fa; font-weight: bold;")
        
        self.start_ai_with_model(prompt, model, mode="jod_two_tier")
    
    def finish_jod_two_tier_processing(self):
        """Finalize after both Jod passes complete"""
        # Find remaining empty rows (not filled by either pass)
        missing_items = []
        for row in self.jod_data_rows:
            parts = row.split('\t')
            if len(parts) >= 1:
                name = parts[0].strip()
                label_th = parts[1].strip() if len(parts) >= 2 else ""
                if name.lower() not in self.jod_first_pass_results:
                    missing_items.append({
                        'Name': name,
                        'Label_TH': label_th
                    })
        
        # If there are missing items, ask user
        if missing_items:
            reply = QMessageBox.question(
                self, 
                "Var Sheet: พบรายการที่ไม่ได้แปล",
                f"พบ {len(missing_items)} รายการที่ AI หาไม่เจอ\n\nต้องการกรอกเองไหม?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                dialog = MissingValuesDialog(missing_items, item_type="jod", parent=self)
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    user_inputs = dialog.get_user_inputs()
                    # Merge user inputs into jod_first_pass_results
                    for name, var_eng in user_inputs.items():
                        self.jod_first_pass_results[name] = var_eng
        
        # Convert results to list format
        self.jod_result_data = []
        for name, var_eng in self.jod_first_pass_results.items():
            self.jod_result_data.append({'Name': name, 'VAR_ENG': var_eng})
        
        if self.is_all_mode:
            # Continue to Code processing
            self.lbl_status.setText("✅ Prompt โจทย์ (2-Tier) เสร็จ → กำลังรัน Code...")
            self.run_prompt_code()
        else:
            self.lbl_status.setText("✅ เสร็จสิ้น (2-Tier)!")
            self.set_buttons_enabled(True)
            self.lbl_status.setStyleSheet("color: #a6e3a1; font-weight: bold;")
        
    def run_prompt_code(self):
        if not self.validate_inputs("code"):
            return
        
        # Reset Session Stats (if not continuing from ALL mode)
        if not self.is_all_mode:
            self.session_total_tokens = 0
            self.session_total_time_ms = 0
        
        # Initialize two-tier processing
        self.two_tier_pass = 1  # Track which pass we're on
        self.first_pass_results = {}  # Store results from first pass
        self.empty_rows_for_second_pass = []  # Rows that need second pass
        
        # Prepare data (NO chunking - send all at once)
        val_lines = self.cat_program_data['value'].strip().split('\n')
        header = val_lines[0]
        data_rows = val_lines[1:]
        
        if not data_rows:
            QMessageBox.warning(self, "Warning", "ไม่พบข้อมูลใน Value Sheet")
            return
        
        # Store for later use
        self.code_header = header
        self.code_data_rows = data_rows
        
        # Start Pass 1 with cheap model
        self.run_code_pass(1)

    def run_code_pass(self, pass_number):
        """Run a single pass of code translation"""
        self.two_tier_pass = pass_number
        
        selected_model = self.model_map.get(self.model_combo.currentText(), self.model_combo.currentText())
        is_lite_mode = (selected_model == "google/gemini-2.5-flash-lite")
        
        if pass_number == 1:
            # Pass 1
            if is_lite_mode:
                model = "google/gemini-2.5-flash-lite"
                status_msg = "⏳ Pass 1/2: พนักงานตัวน้อยกำลังทำงาน (Flash 2.5)..."
            else:
                model = selected_model
                agent_name = self.model_combo.currentText()
                status_msg = f"⏳ Pass 1/2: ใช้ {agent_name}..."

            data_text = self.code_header + '\n' + '\n'.join(self.code_data_rows)
        else:
            # Pass 2
            if not self.empty_rows_for_second_pass:
                # No empty rows, we're done
                self.finish_two_tier_processing()
                return
                
            if is_lite_mode:
                model = "google/gemini-3-flash-preview"
                status_msg = f"⏳ Pass 2/2: CEO กำลังดูงาน (Flash 3) สำหรับ {len(self.empty_rows_for_second_pass)} แถวที่ว่าง..."
            else:
                model = selected_model
                agent_name = self.model_combo.currentText()
                status_msg = f"⏳ Pass 2/2: ใช้ {agent_name} เก็บตก {len(self.empty_rows_for_second_pass)} แถว..."

            data_text = self.code_header + '\n' + '\n'.join(self.empty_rows_for_second_pass)
        
        combined_data = f"{self.questionnaire_data}\n\n{data_text}"
        prompt = PROMPT_CODE.replace("[ *** วางข้อมูล Word และ Excel *** ]", combined_data)
        
        self.current_prompt_type = "code"
        self.lbl_status.setText(status_msg)
        self.lbl_status.setStyleSheet("color: #89b4fa; font-weight: bold;")
        
        # Start AI with specific model
        self.start_ai_with_model(prompt, model, mode="two_tier")

    def start_ai_with_model(self, prompt, model, mode="normal"):
        """Start AI with a specific model (not from dropdown)"""
        self.set_buttons_enabled(False)
        self.output_text.clear()
        self.current_ai_mode = mode
        
        # === SAVE PROMPT LOG ===
        try:
            import datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # กำหนดชื่อไฟล์ตามประเภท prompt
            if self.current_prompt_type == "jod":
                log_filename = f"prompt_log_JOD_{timestamp}.txt"
            elif self.current_prompt_type == "code":
                log_filename = f"prompt_log_CODE_{timestamp}.txt"
            else:
                log_filename = f"prompt_log_{timestamp}.txt"
            
            # บันทึกในโฟลเดอร์เดียวกับไฟล์ SPSS หรือ Desktop
            if hasattr(self, 'spss_filepath') and self.spss_filepath:
                log_dir = os.path.dirname(self.spss_filepath)
            else:
                log_dir = os.path.expanduser("~/Desktop")
            
            log_path = os.path.join(log_dir, log_filename)
            
            with open(log_path, 'w', encoding='utf-8') as f:
                f.write(f"=== PROMPT LOG ({self.current_prompt_type.upper()}) ===\n")
                f.write(f"Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Model: {model}\n")
                f.write(f"Mode: {mode}\n")
                f.write("=" * 60 + "\n\n")
                f.write(prompt)
            
            print(f"✅ Prompt log saved: {log_path}")
        except Exception as e:
            print(f"⚠️ Failed to save prompt log: {e}")
        # === END SAVE PROMPT LOG ===
        
        self.worker = AIWorker(prompt, model)
        self.worker.progress.connect(self.on_ai_progress)
        self.worker.finished.connect(self.on_ai_finished)
        self.worker.start()

    def finish_two_tier_processing(self):
        """Finalize after both passes complete"""
        # Find remaining empty rows (not filled by either pass)
        missing_items = []
        for row in self.code_data_rows:
            parts = row.split('\t')
            if len(parts) >= 3:
                var = parts[0].strip()
                val = parts[1].strip()
                label_th = parts[2].strip() if len(parts) >= 3 else ""
                key = (var.lower(), self._normalize_val(val))
                if key not in self.first_pass_results:
                    missing_items.append({
                        'Variable': var,
                        'Value': val,
                        'Label_TH': label_th
                    })
        
        # If there are missing items, ask user
        if missing_items:
            reply = QMessageBox.question(
                self, 
                "พบรายการที่ไม่ได้แปล",
                f"พบ {len(missing_items)} รายการที่ AI หาไม่เจอ\n\nต้องการกรอกเองไหม?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                dialog = MissingValuesDialog(missing_items, item_type="code", parent=self)
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    user_inputs = dialog.get_user_inputs()
                    # Merge user inputs into first_pass_results
                    for key, label in user_inputs.items():
                        self.first_pass_results[key] = label
        
        # Proceed to save
        code_result_list = []
        for key, label in self.first_pass_results.items():
            code_result_list.append({'Variable': key[0], 'Value': key[1], 'Label_EN': label})
        
        if self.is_all_mode:
            self.is_all_mode = False
            self.save_all_to_excel(self.jod_result_data, code_result_list)
            self.lbl_status.setText("✅ ALL Mode (2-Tier) เสร็จสมบูรณ์!")
        else:
            self.lbl_status.setText("✅ เสร็จสิ้น (2-Tier)!")
        
        self.set_buttons_enabled(True)
        self.lbl_status.setStyleSheet("color: #a6e3a1; font-weight: bold;")

    def start_ai(self, prompt, mode="normal"):
        self.set_buttons_enabled(False)
        if mode == "normal" or (mode == "code_chunk" and self.current_chunk_index == 0):
            self.output_text.clear()
            
        if mode == "normal":
            self.lbl_status.setText("⏳ กำลังประมวลผล...")
            
        self.lbl_status.setStyleSheet("color: #89b4fa; font-weight: bold;")
        self.current_ai_mode = mode # Store mode to handle callback
        
        model = self.model_map.get(self.model_combo.currentText(), self.model_combo.currentText())
        self.worker = AIWorker(prompt, model)
        self.worker.progress.connect(self.on_ai_progress)
        self.worker.finished.connect(self.on_ai_finished)
        self.worker.start()
        
    def run_prompt_all(self):
        if not self.validate_inputs("all"):
            return
        
        # Reset Session Stats
        self.session_total_tokens = 0
        self.session_total_time_ms = 0
        
        self.is_all_mode = True
        self.jod_result_data = []
        self.run_prompt_jod()
        
    def validate_inputs(self, mode):
        if not self.ensure_openrouter_api_key():
            return False
        if not self.questionnaire_data:
            QMessageBox.warning(self, "ข้อมูลไม่ครบ", "กรุณาเลือกไฟล์แบบสอบถามก่อน")
            return False
        if mode in ["jod", "all"] and not self.cat_program_data.get("var"):
            QMessageBox.warning(self, "ข้อมูลไม่ครบ", "กรุณาโหลดไฟล์ SPSS ก่อน (ไม่พบข้อมูล Var)")
            return False
        if mode in ["code", "all"] and not self.cat_program_data.get("value"):
            QMessageBox.warning(self, "ข้อมูลไม่ครบ", "กรุณาโหลดไฟล์ SPSS ก่อน (ไม่พบข้อมูล Value)")
            return False
        return True

    def ensure_openrouter_api_key(self):
        if get_openrouter_api_key():
            return True

        key, ok = QInputDialog.getText(
            self,
            "OpenRouter API Key",
            "กรุณาใส่ Key (ใส่ครั้งเดียว):",
            QLineEdit.EchoMode.Password
        )
        if not ok:
            return False

        key = key.strip()
        if not key:
            QMessageBox.warning(self, "API Key", "กรุณากรอก API Key ให้ถูกต้อง")
            return False

        saved_path = save_openrouter_api_key(key)
        if not saved_path:
            QMessageBox.critical(
                self,
                "API Key",
                "บันทึก openrouter.json ไม่สำเร็จ (สิทธิ์ไม่พอหรือพาธไม่ถูกต้อง)"
            )
            return False

        QMessageBox.information(
            self,
            "API Key",
            f"บันทึก API Key แล้วที่:\n{saved_path}"
        )
        return True
        

        
    def on_ai_progress(self, content):
        if content.startswith('<span') or content.startswith('<i>'):
            self.output_text.insertHtml(content)
        else:
            self.output_text.insertPlainText(content)
            
        self.output_text.verticalScrollBar().setValue(
            self.output_text.verticalScrollBar().maximum()
        )
        
    def on_ai_finished(self, success, result, tokens, time_ms):
        self.ai_result_text = result # Will be overwritten if chunking
        
        # Accumulate stats
        if not hasattr(self, 'session_total_tokens'):
             self.session_total_tokens = 0
             self.session_total_time_ms = 0
        
        if success:
             self.session_total_tokens += tokens
             self.session_total_time_ms += time_ms
        
        self.total_tokens_used = self.session_total_tokens
        
        # Format TOTAL time
        minutes = self.session_total_time_ms // 60000
        seconds = (self.session_total_time_ms % 60000) // 1000
        if minutes > 0:
            time_str = f"{minutes} นาที {seconds} วินาที"
        else:
            time_str = f"{seconds} วินาที"
        
        if success:
            # Handle Jod Two-Tier Processing
            if hasattr(self, 'current_ai_mode') and self.current_ai_mode == "jod_two_tier":
                parsed_results = self.parse_ai_result("jod")
                
                if self.jod_two_tier_pass == 1:
                    # Pass 1 finished - store results and find empties
                    filled_count = 0
                    for item in parsed_results:
                        name = item['Name'].strip().lower()
                        var_eng = item.get('VAR_ENG', '').strip()
                        if var_eng:
                            self.jod_first_pass_results[name] = var_eng
                            filled_count += 1
                    
                    # Find rows that are still empty
                    self.jod_empty_rows_for_second_pass = []
                    for row in self.jod_data_rows:
                        parts = row.split('\t')
                        if len(parts) >= 1:
                            name = parts[0].strip().lower()
                            if name not in self.jod_first_pass_results:
                                self.jod_empty_rows_for_second_pass.append(row)
                    
                    empty_count = len(self.jod_empty_rows_for_second_pass)
                    self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens | Jod Pass 1: {filled_count} เติม, {empty_count} ว่าง")
                    
                    if empty_count > 0:
                        self.lbl_status.setText(f"✅ Jod Pass 1 เสร็จ → เริ่ม Pass 2 ({empty_count} แถว)")
                        self.run_jod_pass(2)
                        return
                    else:
                        self.finish_jod_two_tier_processing()
                        self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens")
                        return
                else:
                    # Pass 2 finished - merge results
                    pass2_filled = 0
                    for item in parsed_results:
                        name = item['Name'].strip().lower()
                        var_eng = item.get('VAR_ENG', '').strip()
                        if var_eng and name not in self.jod_first_pass_results:
                            self.jod_first_pass_results[name] = var_eng
                            pass2_filled += 1
                    
                    self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens | Jod Pass 2: +{pass2_filled} เติมเพิ่ม")
                    self.finish_jod_two_tier_processing()
                    return

            # Handle Code Two-Tier Processing (Value Sheet)
            if hasattr(self, 'current_ai_mode') and self.current_ai_mode == "two_tier":
                parsed_results = self.parse_ai_result("code")
                
                if self.two_tier_pass == 1:
                    # Pass 1 finished - store results and find empties
                    filled_count = 0
                    for item in parsed_results:
                        key = (item['Variable'].strip().lower(), self._normalize_val(item['Value']))
                        label = item.get('Label_EN', '').strip()
                        if label:
                            self.first_pass_results[key] = label
                            filled_count += 1
                    
                    # Find rows that are still empty
                    self.empty_rows_for_second_pass = []
                    for row in self.code_data_rows:
                        parts = row.split('\t')
                        if len(parts) >= 2:
                            key = (parts[0].strip().lower(), self._normalize_val(parts[1]))
                            if key not in self.first_pass_results:
                                self.empty_rows_for_second_pass.append(row)
                    
                    empty_count = len(self.empty_rows_for_second_pass)
                    self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens | Pass 1: {filled_count} เติม, {empty_count} ว่าง")
                    
                    if empty_count > 0:
                        # Proceed to Pass 2
                        self.lbl_status.setText(f"✅ Pass 1 เสร็จ → เริ่ม Pass 2 ({empty_count} แถว)")
                        self.run_code_pass(2)
                        return
                    else:
                        # All done, no empties
                        self.finish_two_tier_processing()
                        self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens")
                        return
                        
                else:
                    # Pass 2 finished - merge results
                    pass2_filled = 0
                    for item in parsed_results:
                        key = (item['Variable'].strip().lower(), self._normalize_val(item['Value']))
                        label = item.get('Label_EN', '').strip()
                        if label and key not in self.first_pass_results:
                            self.first_pass_results[key] = label
                            pass2_filled += 1
                    
                    self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens | Pass 2: +{pass2_filled} เติมเพิ่ม")
                    self.finish_two_tier_processing()
                    return

            # Check ALL mode (JOD finished)
            if self.is_all_mode and self.current_prompt_type == "jod":
                # Save JOD result and continue with Code
                self.jod_result_data = self.parse_ai_result("jod")
                self.lbl_status.setText("✅ Prompt โจทย์ เสร็จ → กำลังรัน Code...")
                self.lbl_stats.setText(f"⏱ {time_str} (Running) | 📊 {self.session_total_tokens:,} tokens")
                
                # Run prompt code (will handle chunking internally)
                self.run_prompt_code() # This initiates chunking
                return
                
            elif self.is_all_mode and self.current_prompt_type == "code":
                 pass

            if not hasattr(self, 'current_ai_mode') or self.current_ai_mode == "normal":
                 if self.is_all_mode:
                      pass
                 else:
                      self.lbl_status.setText("✅ เสร็จสิ้น!")

            self.lbl_status.setStyleSheet("color: #a6e3a1; font-weight: bold;")
            self.lbl_stats.setText(f"⏱ {time_str} | 📊 {self.session_total_tokens:,} tokens")
        else:
            self.is_all_mode = False
            self.lbl_status.setText("❌ เกิดข้อผิดพลาด")
            self.lbl_status.setStyleSheet("color: #f38ba8; font-weight: bold;")
            self.lbl_stats.setText(f"⏱ {time_str}")
            QMessageBox.critical(self, "Error", result)
        
        self.set_buttons_enabled(True)
        
    def set_buttons_enabled(self, enabled):
        self.btn_prompt_jod.setEnabled(enabled)
        self.btn_prompt_code.setEnabled(enabled)
        self.btn_prompt_all.setEnabled(enabled)
        
    # === Parse & Save Methods ===
    def parse_ai_result(self, prompt_type):
        result = []
        lines = self.ai_result_text.strip().split('\n')
        thai_re = re.compile(r"[\u0E00-\u0E7F]")
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('```') or '---' in line:
                continue
            
            if line.startswith('|'):
                cols = [c.strip() for c in line.split('|') if c.strip()]
                if 'Name' in line or 'Variable' in line or 'VAR_ENG' in line or 'Label_EN' in line:
                    continue
            elif '\t' in line:
                cols = line.split('\t')
                if cols[0].lower() in ['name', 'variable']:
                    continue
            else:
                continue
                
            if prompt_type == "jod" and len(cols) >= 2:
                var_eng = cols[1] if len(cols) > 1 else ''
                if var_eng and thai_re.search(var_eng):
                    var_eng = ''
                # === VALIDATION: ตรวจสอบว่า VAR_ENG มีอยู่ในแบบสอบถามจริงหรือไม่ ===
                if var_eng:
                    var_eng = self._validate_english_text(var_eng)
                result.append({'Name': cols[0], 'VAR_ENG': var_eng})
            elif prompt_type == "code" and len(cols) >= 3:
                label_en = cols[2] if len(cols) > 2 else ''
                if label_en and thai_re.search(label_en):
                    label_en = ''
                # === VALIDATION: ตรวจสอบว่า Label_EN มีอยู่ในแบบสอบถามจริงหรือไม่ ===
                if label_en:
                    label_en = self._validate_english_text(label_en)
                result.append({'Variable': cols[0], 'Value': cols[1], 'Label_EN': label_en})
        
        return result
    
    def _validate_english_text(self, eng_text):
        """
        ตรวจสอบว่าข้อความภาษาอังกฤษที่ AI ให้มา มีอยู่ในแบบสอบถามจริงหรือไม่
        ถ้าไม่มี จะ return '' (ค่าว่าง)
        """
        if not eng_text or not hasattr(self, 'questionnaire_data') or not self.questionnaire_data:
            return eng_text  # ถ้าไม่มี questionnaire_data ให้ return ตามเดิม
        
        clean_text = eng_text.strip()
        
        # Normalize: แปลง whitespace หลายอันเป็น space เดียว และ lowercase
        def normalize(text):
            # แปลง en-dash, em-dash เป็น hyphen
            text = text.replace('–', '-').replace('—', '-')
            # ลบ newlines, tabs, multiple spaces
            text = re.sub(r'\s+', ' ', text)
            return text.lower().strip()
        
        questionnaire_normalized = normalize(self.questionnaire_data)
        
        # === ตรวจสอบพิเศษสำหรับ (Rxx) pattern ===
        rxx_match = re.match(r'^\(R\d+\)', clean_text)
        if rxx_match:
            rxx_code = rxx_match.group(0).lower()  # เช่น "(r310)"
            # ถ้า (Rxx) มีอยู่ในแบบสอบถาม -> ผ่าน
            if rxx_code in questionnaire_normalized:
                return eng_text  # พบ (Rxx) ในแบบสอบถาม - ใช้ได้
        
        # ลบ prefix เช่น "S37) " หรือ "Q1A) " ออกก่อนตรวจสอบ
        prefix_match = re.match(r'^[A-Za-z0-9_]+\)\s*', clean_text)
        if prefix_match:
            clean_text = clean_text[prefix_match.end():].strip()
        
        if not clean_text:
            return ''
        
        clean_text_normalized = normalize(clean_text)
        
        # ตรวจสอบแบบตรงตัว (exact substring)
        if clean_text_normalized in questionnaire_normalized:
            return eng_text  # พบในแบบสอบถาม - ใช้ได้
        
        # ตรวจสอบ 3 คำแรก (ผ่อนคลายจาก 5 คำ)
        words = clean_text_normalized.split()
        if len(words) >= 3:
            first_3_words = ' '.join(words[:3])
            if first_3_words in questionnaire_normalized:
                return eng_text  # พบบางส่วนในแบบสอบถาม - ใช้ได้
        
        # ตรวจสอบ 2 คำแรก (ผ่อนคลายมากขึ้นสำหรับข้อความสั้น)
        if len(words) >= 2:
            first_2_words = ' '.join(words[:2])
            if first_2_words in questionnaire_normalized:
                return eng_text  # พบบางส่วนในแบบสอบถาม - ใช้ได้
        
        # ไม่พบในแบบสอบถาม - AI อาจแปลเอง - return ค่าว่าง
        print(f"⚠️ VALIDATION: ไม่พบข้อความนี้ในแบบสอบถาม (อาจเป็นการแปลเอง): {eng_text[:80]}...")
        return ''
        
    def download_excel(self):
        if not self.ai_result_text:
            QMessageBox.warning(self, "ไม่มีผลลัพธ์", "กรุณารัน AI ก่อน")
            return
        if not self.cat_program_filepath:
            QMessageBox.warning(self, "ไม่มีไฟล์", "กรุณาเลือกไฟล์โปรแกรมแมวก่อน")
            return
            
        parsed_data = self.parse_ai_result(self.current_prompt_type)
        self.save_single_sheet(parsed_data)
        
    def save_single_sheet(self, data):
        try:
            wb = openpyxl.load_workbook(self.cat_program_filepath)
            sheet_names_lower = [s.lower() for s in wb.sheetnames]
            updated = 0
            
            if self.current_prompt_type == "jod":
                if "var" not in sheet_names_lower:
                    return
                sheet = wb[wb.sheetnames[sheet_names_lower.index("var")]]
                name_col, eng_col = self.find_columns(sheet, ['name'], 'var_eng', 'VAR_ENG')
                if name_col:
                    # Use lowercase keys for case-insensitive matching
                    lookup = {item['Name'].lower(): item['VAR_ENG'] for item in data}
                    for row in range(2, sheet.max_row + 1):
                        name = str(sheet.cell(row, name_col).value or '').strip()
                        name_lower = name.lower()
                        if name_lower in lookup and lookup[name_lower]:
                            sheet.cell(row, eng_col, lookup[name_lower])
                            updated += 1
            else:
                if "value" not in sheet_names_lower:
                    return
                sheet = wb[wb.sheetnames[sheet_names_lower.index("value")]]
                var_col, val_col, eng_col = self.find_columns_code(sheet)
                if var_col and val_col:

                    # Normalize keys in lookup
                    lookup = {}
                    for item in data:
                        norm_var = item['Variable'].strip().lower()
                        norm_val = self._normalize_val(item['Value'])
                        lookup[(norm_var, norm_val)] = item['Label_EN']

                    for row in range(2, sheet.max_row + 1):
                        var = str(sheet.cell(row, var_col).value or '').strip().lower()
                        val = self._normalize_val(sheet.cell(row, val_col).value or '')
                        
                        if (var, val) in lookup and lookup[(var, val)]:
                            sheet.cell(row, eng_col, lookup[(var, val)])
                            updated += 1
            
            # Save
            name = os.path.splitext(os.path.basename(self.cat_program_filepath))[0]
            save_path, _ = QFileDialog.getSaveFileName(self, "บันทึก Excel", f"{name}_EN.xlsx", "Excel Files (*.xlsx)")
            if save_path:
                wb.save(save_path)
                QMessageBox.information(self, "สำเร็จ", f"บันทึกเรียบร้อย!\nอัพเดท {updated} แถว")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))
            
    def _normalize_val(self, val):
        s = str(val).strip()
        try:
            f = float(s)
            if f.is_integer():
                return str(int(f))
            return str(f)
        except ValueError:
            return s.lower()

    def save_all_to_excel(self, jod_data, code_data):
        try:
            wb = openpyxl.load_workbook(self.cat_program_filepath)
            sheet_names_lower = [s.lower() for s in wb.sheetnames]
            updated_jod = updated_code = 0
            
            # Update Var sheet
            if "var" in sheet_names_lower and jod_data:
                sheet = wb[wb.sheetnames[sheet_names_lower.index("var")]]
                name_col, eng_col = self.find_columns(sheet, ['name'], 'var_eng', 'VAR_ENG')
                if name_col:
                    # Use lowercase keys for case-insensitive matching
                    lookup = {item['Name'].lower(): item['VAR_ENG'] for item in jod_data}
                    for row in range(2, sheet.max_row + 1):
                        name = str(sheet.cell(row, name_col).value or '').strip()
                        name_lower = name.lower()
                        if name_lower in lookup and lookup[name_lower]:
                            sheet.cell(row, eng_col, lookup[name_lower])
                            updated_jod += 1
            
            # Update Value sheet
            if "value" in sheet_names_lower and code_data:
                sheet = wb[wb.sheetnames[sheet_names_lower.index("value")]]
                var_col, val_col, eng_col = self.find_columns_code(sheet)
                if var_col and val_col:
                    # Normalize keys in lookup
                    lookup = {}
                    for item in code_data:
                        norm_var = item['Variable'].strip().lower()
                        norm_val = self._normalize_val(item['Value'])
                        lookup[(norm_var, norm_val)] = item['Label_EN']
                    
                    for row in range(2, sheet.max_row + 1):
                        var = str(sheet.cell(row, var_col).value or '').strip().lower()
                        val = self._normalize_val(sheet.cell(row, val_col).value or '')
                        
                        if (var, val) in lookup and lookup[(var, val)]:
                            sheet.cell(row, eng_col, lookup[(var, val)])
                            updated_code += 1
            
            # Save
            # Use original SPSS filename if available, else use temp/Excel filename
            if hasattr(self, 'lbl_spss_file') and self.lbl_spss_file.text() != "ยังไม่ได้เลือกไฟล์":
                name = os.path.splitext(self.lbl_spss_file.text())[0]
            else:
                name = os.path.splitext(os.path.basename(self.cat_program_filepath))[0]
                
            save_path, _ = QFileDialog.getSaveFileName(self, "บันทึก Excel", f"{name}_ALL_EN.xlsx", "Excel Files (*.xlsx)")
            if save_path:
                wb.save(save_path)
                QMessageBox.information(self, "ALL Mode สำเร็จ!", 
                    f"บันทึกเรียบร้อย!\n\n"
                    f"📝 Var Sheet: {updated_jod} แถว\n"
                    f"📝 Value Sheet: {updated_code} แถว")
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))
            
    def find_columns(self, sheet, search_cols, eng_col_name, eng_col_header):
        name_col = eng_col = None
        for col in range(1, sheet.max_column + 1):
            val = str(sheet.cell(1, col).value or '').strip().lower()
            if val in search_cols:
                name_col = col
            elif val == eng_col_name:
                eng_col = col
        if eng_col is None:
            eng_col = sheet.max_column + 1
            sheet.cell(1, eng_col, eng_col_header)
        return name_col, eng_col
        
    def find_columns_code(self, sheet):
        var_col = val_col = eng_col = None
        for col in range(1, sheet.max_column + 1):
            val = str(sheet.cell(1, col).value or '').strip().lower()
            if val == 'variable':
                var_col = col
            elif val == 'value':
                val_col = col
            elif val == 'label_en':
                eng_col = col
        if eng_col is None:
            eng_col = sheet.max_column + 1
            sheet.cell(1, eng_col, "Label_EN")
        return var_col, val_col, eng_col
        
    # === Utility Methods ===
    def copy_to_clipboard(self):
        text = self.output_text.toPlainText().strip()
        if text:
            pyperclip.copy(text)
            QMessageBox.information(self, "สำเร็จ", "คัดลอกแล้ว!")
        else:
            QMessageBox.warning(self, "ไม่มีข้อความ", "ไม่มีข้อความให้คัดลอก")

    def load_ai_result_excel(self):
        """Load Excel file with AI translation results for SPSS save"""
        filepath, _ = QFileDialog.getOpenFileName(
            self, "เลือกไฟล์ Excel ผลลัพธ์",
            "", "Excel Files (*.xlsx *.xls);;All Files (*.*)"
        )
        if not filepath:
            return
            
        try:
            xls = pd.ExcelFile(filepath, engine="openpyxl")
            loaded_something = False
            
            self.ai_result_var_df = None
            self.ai_result_value_df = None
            
            def _normalize(df):
                df.columns = [str(c).strip() for c in df.columns]
                for c in df.columns:
                    if df[c].dtype == object:
                        df[c] = df[c].astype(str).str.strip()
                return df
            
            # Var sheet
            if 'Var' in xls.sheet_names:
                self.ai_result_var_df = pd.read_excel(
                    xls, sheet_name='Var',
                    dtype=str, keep_default_na=False, na_filter=False
                )
                self.ai_result_var_df = _normalize(self.ai_result_var_df)
                loaded_something = True
                
            # Value sheet
            if 'Value' in xls.sheet_names:
                self.ai_result_value_df = pd.read_excel(
                    xls, sheet_name='Value',
                    dtype=str, keep_default_na=False, na_filter=False
                )
                self.ai_result_value_df = _normalize(self.ai_result_value_df)
                
                # Normalize Label_EN column name
                for cand in ['Label_EN', 'Label_En', 'label_en', 'Label EN', 'LABEL_EN']:
                    if cand in self.ai_result_value_df.columns:
                        if cand != 'Label_EN':
                            self.ai_result_value_df.rename(columns={cand: 'Label_EN'}, inplace=True)
                        break
                loaded_something = True
            
            if loaded_something:
                self.btn_save_spss.setEnabled(True)
                var_count = len(self.ai_result_var_df) if self.ai_result_var_df is not None else 0
                val_count = len(self.ai_result_value_df) if self.ai_result_value_df is not None else 0
                QMessageBox.information(
                    self, "โหลดสำเร็จ",
                    f"โหลดข้อมูลจาก Excel เรียบร้อย!\n\n"
                    f"📊 Var: {var_count} แถว\n"
                    f"📊 Value: {val_count} แถว\n\n"
                    f"กด 'Save to SPSS' เพื่อบันทึกลงไฟล์ SPSS"
                )
            else:
                QMessageBox.warning(self, "ไม่พบข้อมูล", 
                    "ไม่พบชีทชื่อ 'Var' หรือ 'Value' ในไฟล์ Excel ที่เลือก")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"ไม่สามารถอ่านไฟล์ Excel ได้:\n{e}")

    def save_labels_to_spss(self):
        """Save translated labels back to SPSS file using pyspssio"""
        try:
            import pyspssio
        except ImportError:
            QMessageBox.critical(
                self,
                "ไลบรารีไม่พร้อม",
                "ไม่พบโมดูล 'pyspssio'\n\nติดตั้งด้วยคำสั่ง:\n  pip install pyspssio"
            )
            return

        if not hasattr(self, 'spss_filepath') or not self.spss_filepath:
            QMessageBox.warning(self, "ไม่มีไฟล์ SPSS", "กรุณาโหลดไฟล์ SPSS ก่อน")
            return

        # Check if we have Excel data loaded
        has_excel_var = hasattr(self, 'ai_result_var_df') and self.ai_result_var_df is not None
        has_excel_value = hasattr(self, 'ai_result_value_df') and self.ai_result_value_df is not None
        
        if not has_excel_var and not has_excel_value:
            QMessageBox.warning(self, "ไม่มีข้อมูล", 
                "กรุณากด 'โหลดไฟล์ Excel' เพื่อเลือกไฟล์ที่ต้องการบันทึกลง SPSS")
            return

        # Helper: find column case-insensitive
        def find_col(df, names):
            for name in names:
                for col in df.columns:
                    if str(col).lower() == str(name).lower():
                        return col
            return None

        # Helper: coerce value type to match existing keys
        def _coerce_value_type(val, var, value_labels, meta):
            if var in value_labels and len(value_labels[var]) > 0:
                sample_key = next(iter(value_labels[var].keys()))
                try:
                    if isinstance(sample_key, int):
                        return int(float(val))
                    elif isinstance(sample_key, float):
                        return float(val)
                    else:
                        return str(val)
                except Exception:
                    return sample_key.__class__(val)
            try:
                is_numeric = (self.spss_meta.variable_types.get(var, 0) == 0)
            except Exception:
                is_numeric = True
            if is_numeric:
                f = float(val)
                return int(f) if abs(f - int(f)) < 1e-9 else f
            return str(val)

        try:
            meta = self.spss_meta
            original_columns = list(meta.column_names)
            df_to_save = self.original_spss_df[original_columns].copy()

            # Copy existing labels
            new_var_labels_list = list(meta.column_labels[:])
            new_value_labels = {k: dict(v) for k, v in meta.variable_value_labels.items()}

            delimiter = self.ma_combo.currentText()
            import re

            # ====== Update Value Labels from Excel ======
            if has_excel_value:
                lbl_en_col = find_col(self.ai_result_value_df, ["Label_EN", "Label_En", "label_en"])
                if lbl_en_col:
                    for _, r in self.ai_result_value_df.iterrows():
                        var = r.get("Variable")
                        val = r.get("Value")
                        en_raw = r.get(lbl_en_col)
                        
                        if not var or pd.isna(var):
                            continue
                        if pd.isna(val):
                            continue
                        
                        # Find original variable name
                        original_var = None
                        for col in original_columns:
                            if col.lower() == str(var).lower():
                                original_var = col
                                break
                        if not original_var:
                            continue
                        
                        en = "" if pd.isna(en_raw) else str(en_raw).strip()
                        if not en:
                            continue
                        
                        try:
                            vv = _coerce_value_type(val, original_var, new_value_labels, meta)
                        except Exception:
                            continue
                        
                        new_value_labels.setdefault(original_var, {})
                        new_value_labels[original_var][vv] = en

            # Propagate MA value labels (_O1, _O2, ...) to same group
            if delimiter:
                vars_with_ma = [v for v in df_to_save.columns
                                if re.search(rf'{re.escape(delimiter)}\d+$', v)]
                groups = {}
                for v in vars_with_ma:
                    base = v.rsplit(delimiter, 1)[0]
                    groups.setdefault(base, []).append(v)

                for base, members in groups.items():
                    prefer_key = f"{base}{delimiter}1"
                    src_key = None
                    if prefer_key in new_value_labels and new_value_labels[prefer_key]:
                        src_key = prefer_key
                    else:
                        for m in members:
                            if m in new_value_labels and new_value_labels[m]:
                                src_key = m
                                break
                    if not src_key:
                        continue
                    src = dict(new_value_labels[src_key])
                    for m in members:
                        new_value_labels[m] = dict(src)

            # ====== Update Variable Labels from Excel ======
            if has_excel_var:
                eng_col = find_col(self.ai_result_var_df, ["VAR_ENG", "Var_Eng", "var_eng"])
                if eng_col:
                    for _, r in self.ai_result_var_df.iterrows():
                        name = r.get("Name")
                        eng = r.get(eng_col)
                        if pd.isna(name) or pd.isna(eng):
                            continue
                        name_str = str(name).strip()
                        eng_str = str(eng).strip()
                        if not eng_str:
                            continue
                        # Find and update
                        for i, col in enumerate(original_columns):
                            if col.lower() == name_str.lower():
                                new_var_labels_list[i] = eng_str
                                break

            # Propagate MA var labels to same group
            if delimiter:
                bases = {v.rsplit(delimiter, 1)[0]
                        for v in df_to_save.columns
                        if re.search(rf'{re.escape(delimiter)}\d+$', v)}
                for base in bases:
                    key = f"{base}{delimiter}1"
                    if key in original_columns:
                        idx = original_columns.index(key)
                        lbl = new_var_labels_list[idx]
                        if pd.notna(lbl) and lbl:
                            for v in [vv for vv in df_to_save.columns
                                    if re.search(rf'^{re.escape(base)}{re.escape(delimiter)}\d+$', vv)]:
                                new_var_labels_list[original_columns.index(v)] = lbl

            # Prepare metadata for pyspssio
            var_labels = {c: ("" if pd.isna(l) else str(l))
                        for c, l in zip(original_columns, new_var_labels_list)}
            meta_out = {"var_labels": var_labels, "var_value_labels": new_value_labels}

            # Ask for save location
            base_name = os.path.splitext(os.path.basename(self.spss_filepath))[0]
            save_path, _ = QFileDialog.getSaveFileName(
                self, "บันทึกไฟล์ SPSS ใหม่",
                f"{base_name}_EN.sav",
                "SPSS Files (*.sav)"
            )
            if not save_path:
                return

            pyspssio.write_sav(save_path, df_to_save, metadata=meta_out)
            QMessageBox.information(
                self, "บันทึกสำเร็จ",
                f"ไฟล์ SPSS ถูกบันทึกเรียบร้อยแล้ว!\n\n{save_path}"
            )
            
            # Clear loaded data
            self.ai_result_var_df = None
            self.ai_result_value_df = None

        except Exception as e:
            QMessageBox.critical(self, "บันทึกไม่สำเร็จ", f"เกิดข้อผิดพลาด:\n{e}")


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    app.setStyleSheet(DARK_STYLE)
    
    window = GenPromtApp()
    window.show()
    
    sys.exit(app.exec())


#if __name__ == "__main__":
    #main()


# <<< START OF CHANGES >>>
# --- ฟังก์ชัน Entry Point ใหม่ (สำหรับให้ Launcher เรียก) ---
def run_this_app(working_dir=None): # ชื่อฟังก์ชันนี้จะถูกใช้ใน Launcher
    """
    ฟังก์ชันหลักสำหรับสร้างและรัน QuotaSamplerApp.
    """
    print(f"--- QUOTA_SAMPLER_INFO: Starting 'QuotaSamplerApp' via run_this_app() ---")
    try:
        # --- โค้ดที่ย้ายมาจาก if __name__ == "__main__": เดิมจะมาอยู่ที่นี่ ---
    #if __name__ == "__main__":
        main()  


        print(f"--- QUOTA_SAMPLER_INFO: QuotaSamplerApp mainloop finished. ---")

    except Exception as e:
        # ดักจับ Error ที่อาจเกิดขึ้นระหว่างการสร้างหรือรัน App
        print(f"QUOTA_SAMPLER_ERROR: An error occurred during QuotaSamplerApp execution: {e}")
        # แสดง Popup ถ้ามีปัญหา
        if 'root' not in locals() or not root.winfo_exists(): # สร้าง root ชั่วคราวถ้ายังไม่มี
            root_temp = tk.Tk()
            root_temp.withdraw()
            messagebox.showerror("Application Error (Quota Sampler)",
                               f"An unexpected error occurred:\n{e}", parent=root_temp)
            root_temp.destroy()
        else:
            messagebox.showerror("Application Error (Quota Sampler)",
                               f"An unexpected error occurred:\n{e}", parent=root) # ใช้ root ที่มีอยู่ถ้าเป็นไปได้
        sys.exit(f"Error running QuotaSamplerApp: {e}") # อาจจะ exit หรือไม่ก็ได้ ขึ้นกับการออกแบบ


# --- ส่วน Run Application เมื่อรันไฟล์นี้โดยตรง (สำหรับ Test) ---
if __name__ == "__main__":
    print("--- Running QuotaSamplerApp.py directly for testing ---")
    # (ถ้ามีการตั้งค่า DPI ด้านบน มันจะทำงานอัตโนมัติ)

    # เรียกฟังก์ชัน Entry Point ที่เราสร้างขึ้น
    run_this_app()

    print("--- Finished direct execution of QuotaSamplerApp.py ---")
# <<< END OF CHANGES >>>
