import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog, StringVar
from tkinter import ttk
from tkinter import font as tkFont
import os
import pyperclip
import webbrowser
import io

# --- ส่วนของการอ่านไฟล์ (จำเป็นสำหรับโปรแกรมนี้) ---
try:
    import docx
    import openpyxl
    import pandas as pd
    import pyreadstat
    import pyspssio
    LIBS_INSTALLED = True
except ImportError:
    LIBS_INSTALLED = False

# --- แม่แบบ Prompt (Templates) ---
# ... (คง PROMPT_JOD และ PROMPT_CODE ไว้เหมือนเดิม) ...
PROMPT_JOD = """คุยไทยนะ
จากข้อมูลต่อไปนี้:

[ *** วางข้อมูล Word และ Excel *** ]

ให้ดำเนินการดังนี้:

1.  **ระบุคู่ข้อมูล:** สำหรับแต่ละแถวในตารางข้อมูลสุดท้าย (Name, VAR_THA):
    *   ค้นหาข้อความภาษาไทย (VAR_THA) ที่ตรงกันในส่วนข้อมูลต้นฉบับ
    *   ถ้าใน (VAR_THA) มีคำว่า Hide ไม่ต้องสนใจให้ข้ามบรรทัดนั้นไปแต่ให้คงตำแหน่งข้อความใน (VAR_THA) ไว้เหมือนเดิม
    *   มองหาข้อความภาษาอังกฤษ**ที่ปรากฏคู่กันโดยตรง**กับ VAR_THA นั้นๆ ในข้อมูลต้นฉบับ (เช่น อยู่ในคอลัมน์ติดกัน, อยู่ในบรรทัดเดียวกันในตำแหน่งที่สอดคล้องกัน, อยู่ในลำดับถัดกัน (ก่อนหน้า / หลัง) ภายในกลุ่มเดียวกัน, หรือมีการระบุว่าเป็นคำแปลอย่างชัดเจน)

2.  **สร้างตารางผลลัพธ์:** สร้างตารางที่มี 3 คอลัมน์ ตามลำดับดังนี้: Name, VAR_THA, VAR_ENG

3.  **เติมข้อมูล:**
    *   นำข้อมูล Name และ VAR_THA จากตารางสุดท้ายมาใส่ในตารางผลลัพธ์
    *   **เฉพาะกรณีที่** พบข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**ตามเงื่อนไขในข้อ 1 เท่านั้น:
        *   **การสร้างส่วนนำหน้า (prefix) สำหรับ VAR_ENG:**
            *   ให้สกัดหมายเลขคำถาม (เช่น S1, Q1, MQ1, Q2.1) จากส่วนเริ่มต้นของข้อมูลในคอลัมน์ Name (เช่น ถ้า Name คือ s1 ให้ใช้ "S1", ถ้า Name คือ q1_1_O1 ให้ใช้ "Q1.1", ถ้า Name คือ mq5 ให้ใช้ "MQ5") หรือจากส่วนเริ่มต้นของ VAR_THA หากมีการระบุหมายเลขคำถามไว้อยอย่างชัดเจน (เช่น "S1 กรุณาระบุ...", "Q1.1 ...")
            *   นำหมายเลขคำถามที่สกัดได้ตามด้วยเครื่องหมายจุด (.) เพื่อสร้างเป็นส่วนนำหน้า (prefix)
        *   นำส่วนนำหน้าที่สร้างขึ้นนี้ (หมายเลขคำถาม + จุด) ต่อด้วยข้อความภาษาอังกฤษที่จับคู่ได้ มาเติมในคอลัมน์ VAR_ENG ของแถวที่สอดคล้องกัน
        *   ตัวอย่างเช่น:
            *   ถ้า Name คือ mq5, VAR_THA คือ "MQ5 คุณรู้สึกอยากซื้อสินค้า...", และข้อความอังกฤษที่คู่กันคือ "To what extent would you want to purchase the product..." ผลลัพธ์ใน VAR_ENG ควรเป็น "MQ5.To what extent would you want to purchase the product..."
            *   ถ้า Name คือ I_1_q1, VAR_THA คือ "(บิสกิต (รวมถึง ชนิดแท่ง)) Q1 กรุณาระบุความถี่...", และข้อความอังกฤษที่คู่กันคือ "Biscuit" ผลลัพธ์ใน VAR_ENG ควรเป็น "Q1.Biscuit" (โดย "Q1." มาจากการตีความ I_1_q1 หรือ "Q1" ใน VAR_THA)
            *   ถ้า Name คือ s1, VAR_THA คือ "S1 กรุณาระบุเพศ...", และข้อความอังกฤษที่คู่กันคือ "Please indicate your gender..." ผลลัพธ์ใน VAR_ENG ควรเป็น "S1.Please indicate your gender..."
    *   **หากไม่พบ**ข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**กับ VAR_THA ในตำแหน่งที่สอดคล้องกันในข้อมูลต้นฉบับ ให้**เว้นว่าง**คอลัมน์ VAR_ENG สำหรับแถวนั้นไว้ **ห้าม**นำข้อมูลภาษาอังกฤษจากส่วนอื่นของเอกสารที่ไม่ใช่คู่โดยตรงมาเติม หรือทำการแปลเอง
    *   ตัวอย่างรูปแบบที่ต้องการสำหรับแถวที่มีการเติมเลขข้อ (เป็นการยืนยันรูปแบบ ไม่ใช่ให้ทำเฉพาะ P16):
        P16 คุณรู้สึกไม่ชอบผลิตภัณฑ์ที่ได้ดื่มไปเมื่อสักครู่นี้ในแง่ใดบ้าง กรุณาตอบทุกข้อที่ไม่ชอบ P16.Regarding the product you drank earlier, please select all the aspects you disliked from the following options.

4.  **จัดรูปแบบ VAR_ENG:**
    *   สำหรับ VAR_ENG ที่มีการเติมข้อมูล: ตรวจสอบให้แน่ใจว่าข้อความ**ไม่มีช่องว่างนำหน้า (leading space)** และ **ไม่มีช่องว่างตามหลัง (trailing space)** (ยกเว้นช่องว่างที่อาจมีระหว่าง Prefix ที่สร้างกับตัวข้อความภาษาอังกฤษ ถ้าการจับคู่เดิมมีช่องว่างนั้นอยู่)
    *   **คง**ช่องว่างที่อยู่**ภายใน**ข้อความภาษาอังกฤษไว้ตามเดิม

5.  **จัดรูปแบบตาราง:**
    *   แสดงผลลัพธ์เป็น **ข้อความธรรมดา (Plain Text)**
    *   ใช้ **ตัวคั่นแท็บ (Tab character)** เพียงตัวเดียวในการแยกระหว่างข้อมูลในแต่ละคอลัมน์
    *   **ห้าม**รวมเซลล์ (No Merging)
"""
PROMPT_CODE = """คุยไทยนะ
จากข้อมูลต่อไปนี้:

[ *** วางข้อมูล Word และ Excel *** ]

ให้ดำเนินการดังนี้:

1.  **ระบุคู่ข้อมูล:** สำหรับแต่ละแถวในตารางข้อมูลสุดท้าย (Variable, Value, Label_Th):
    *   ค้นหาข้อความภาษาไทย (Label_Th) ที่ตรงกันในส่วนข้อมูลต้นฉบับ
    *   ถ้าใน (Label_Th) มีคำว่า Hide ไม่ต้องสนใจให้ข้ามบรรทัดนั้นไปแต่ให้คงตำแหน่งข้อความใน (Label_Th) ไว้เหมือนเดิม
    *   มองหาข้อความภาษาอังกฤษ**ที่ปรากฏคู่กันโดยตรง**กับ Label_Th นั้นๆ ในข้อมูลต้นฉบับ (เช่น อยู่ในคอลัมน์ติดกัน, อยู่ในบรรทัดเดียวกันในตำแหน่งที่สอดคล้องกัน อยู่ในลำดับถัดกัน (ก่อนหน้า / หลัง) ภายในกลุ่มเดียวกัน ,หรือมีการระบุว่าเป็นคำแปลอย่างชัดเจน)
2.  **สร้างตารางผลลัพธ์:** สร้างตารางที่มี 4 คอลัมน์ ตามลำดับดังนี้: Variable, Value, Label_Th, Label_En
3.  **เติมข้อมูล:**
    *   นำข้อมูล Variable, Value, และ Label_Th จากตารางสุดท้ายมาใส่ในตารางผลลัพธ์
    *   **เฉพาะกรณีที่** พบข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**ตามเงื่อนไขในข้อ 1 เท่านั้น ให้นำข้อความภาษาอังกฤษนั้นมาเติมในคอลัมน์ Label_En ของแถวที่สอดคล้องกัน
    *   **หากไม่พบ**ข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**กับ Label_Th ในตำแหน่งที่สอดคล้องกันในข้อมูลต้นฉบับ ให้**เว้นว่าง**คอลัมน์ Label_En สำหรับแถวนั้นไว้ **ห้าม**นำข้อมูลภาษาอังกฤษจากส่วนอื่นของเอกสารที่ไม่ใช่คู่โดยตรงมาเติม หรือทำการแปลเอง
4.  **จัดรูปแบบ Label_En:**
    *   สำหรับ Label_En ที่มีการเติมข้อมูล: ตรวจสอบให้แน่ใจว่าข้อความ**ไม่มีช่องว่างนำหน้า (leading space)** และ **ไม่มีช่องว่างตามหลัง (trailing space)**
    *   **คง**ช่องว่างที่อยู่**ภายใน**ข้อความภาษาอังกฤษไว้ตามเดิม
	ถ้ามี (R1) (R2) หน้า Code ให้เอามาด้วย เช่น
	Variable	Value	Label	Label_EN
	s16a_O1	1	(R1) โซฟี ขอบปกป้อง ผิวสัมผัสแห้ง สลิม มีปีก	(R1) Sofy Side Gather slim (wing)
	s16a_O1	2	(R2) โซฟี ขอบปกป้อง ผิวสัมผัสนุ่ม สลิม มีปีก	(R2) Sofy Side Gather slim Non-mesh (wing)
5.  **จัดรูปแบบตาราง:**
    *   แสดงผลลัพธ์เป็น **ข้อความธรรมดา (Plain Text)**
    *   ใช้ **ตัวคั่นแท็บ (Tab character)** เพียงตัวเดียวในการแยกระหว่างข้อมูลในแต่ละคอลัมน์
    *   **ห้าม**รวมเซลล์ (No Merging)
"""

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("โปรแกรม GetValue+Promt แปะ Eng")
        
        # --- ค่าสีและฟอนต์สำหรับ DARK MODE ---
        self.BG_COLOR = "#252526"
        self.FRAME_BG = "#2d2d30"
        self.TEXT_COLOR = "#cccccc"
        self.TEXT_HEADER_COLOR = "#ffffff"
        self.PRIMARY_COLOR = "#007acc"
        self.SUCCESS_COLOR = "#4CAF50"
        self.ACCENT_COLOR = "#ce9178"
        self.DESTRUCTIVE_COLOR = "#f44747"
        self.PURPLE_COLOR = "#6a0dad"
        self.INPUT_BG = "#3c3c3c"
        self.INPUT_FG = "#f0f0f0"

        self.font_normal = tkFont.Font(family="Tahoma", size=10)
        self.font_bold = tkFont.Font(family="Tahoma", size=11, weight="bold")
        self.font_h1 = tkFont.Font(family="Tahoma", size=13, weight="bold")

        self.root.configure(bg=self.BG_COLOR)
        
        self.style = ttk.Style(self.root)
        self.style.theme_use('clam')
        
        # --- จัดหน้าต่างกลางจอ ---
        window_width = 800
        window_height = 750
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        center_x = int(screen_width/2 - window_width / 2)
        center_y = int(screen_height/2 - window_height / 2)
        self.root.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')
        self.root.minsize(800, 750)

        # --- ตัวแปรสำหรับเก็บข้อมูล ---
        self.questionnaire_data = ""
        self.spss_filepath = None
        self.original_spss_df = None
        self.spss_meta = None
        self.spss_df_var = None
        self.spss_df_value = None
        self.ai_result_var_df = None
        self.ai_result_value_df = None
        self.last_prompt_type = None

        if not LIBS_INSTALLED:
            messagebox.showerror("ไลบรารีไม่ครบ", 
                                 "ยังไม่ได้ติดตั้งไลบรารีที่จำเป็น:\n"
                                 "- python-docx\n- openpyxl\n- pandas\n- pyreadstat\n- pyspssio\n\n"
                                 "โปรแกรมอาจทำงานไม่สมบูรณ์\n"
                                 "กรุณาติดตั้งด้วยคำสั่ง:\n"
                                 "pip install python-docx openpyxl pandas pyreadstat pyspssio")

        self._create_widgets()


    def _create_widgets(self):
        """สร้างและจัดวาง Widget ทั้งหมดในธีม Dark Mode"""
        # --- Style Configurations for Dark Mode ---
        self.style.configure('TFrame', background=self.BG_COLOR)
        self.style.configure('Content.TFrame', background=self.FRAME_BG)
        self.style.configure('TButton', font=self.font_bold, padding=(10, 8), borderwidth=0)
        self.style.configure('TLabel', background=self.FRAME_BG, font=self.font_normal, foreground=self.TEXT_COLOR)
        self.style.configure('Header.TLabel', background=self.BG_COLOR, font=self.font_h1, foreground=self.TEXT_HEADER_COLOR)
        self.style.configure('SubHeader.TLabel', background=self.FRAME_BG, font=self.font_bold, foreground=self.TEXT_HEADER_COLOR)
        self.style.configure('Status.TLabel', background=self.FRAME_BG, font=self.font_normal)
        self.style.configure('TRadiobutton', background=self.FRAME_BG, font=self.font_normal, foreground=self.TEXT_COLOR, indicatorcolor='black')
        self.style.map('TRadiobutton',
            background=[('active', self.FRAME_BG)],
            foreground=[('active', self.TEXT_HEADER_COLOR)]
        )
        
        # --- Custom Button Styles ---
        self.style.configure('Primary.TButton', background=self.PRIMARY_COLOR, foreground='white')
        self.style.map('Primary.TButton', background=[('active', '#0061a1')])
        self.style.configure('Success.TButton', background=self.SUCCESS_COLOR, foreground='white')
        self.style.map('Success.TButton', background=[('active', '#3e8e41')])
        self.style.configure('Accent.TButton', background=self.ACCENT_COLOR, foreground='white')
        self.style.map('Accent.TButton', background=[('active', '#b3765f')])
        self.style.configure('Destructive.TButton', background=self.DESTRUCTIVE_COLOR, foreground='white')
        self.style.map('Destructive.TButton', background=[('active', '#c33b3b')])
        self.style.configure('Purple.TButton', background=self.PURPLE_COLOR, foreground='white')
        self.style.map('Purple.TButton', background=[('active', '#520a8c')])
        
        # --- Main container ---
        main_frame = ttk.Frame(self.root, padding=(20, 15))
        main_frame.pack(fill="both", expand=True)
        main_frame.columnconfigure(0, weight=1)

        # --- 1. ส่วนของการโหลดไฟล์ (Input) ---
        ttk.Label(main_frame, text="1. โหลดไฟล์", style='Header.TLabel').pack(fill="x", anchor="w", pady=(0, 5))
        input_container = ttk.Frame(main_frame, padding=20, style='Content.TFrame')
        input_container.pack(fill="x", pady=(0, 20))
        # ### การแก้ไข Layout ###
        # กำหนดให้คอลัมน์ 2 (ขวาสุด) ยืดขยายได้ ส่วนคอลัมน์อื่นขนาดคงที่
        input_container.columnconfigure(0, weight=0) 
        input_container.columnconfigure(1, weight=0)
        input_container.columnconfigure(2, weight=1) 
        
        # SPSS File
        ttk.Label(input_container, text="ไฟล์ SPSS (.sav):", style='SubHeader.TLabel').grid(row=0, column=0, sticky="e", padx=(0, 10))
        btn_load_spss = ttk.Button(input_container, text="เลือกไฟล์ SPSS", command=self.load_spss_file, style='Primary.TButton')
        btn_load_spss.grid(row=0, column=1, sticky="ew")
        self.label_spss_file = ttk.Label(input_container, text="ยังไม่ได้เลือกไฟล์", foreground="grey", style='Status.TLabel', wraplength=400) # wraplength อาจต้องปรับตามความเหมาะสม
        self.label_spss_file.grid(row=0, column=2, sticky="ew", padx=(15, 0))

        # MA Options
        ma_frame = ttk.Frame(input_container, style='Content.TFrame')
        ma_frame.grid(row=1, column=1, sticky="w", pady=5)
        self.ma_var = StringVar(value="_O")
        ttk.Radiobutton(ma_frame, text="MA=_O", variable=self.ma_var, value="_O").pack(side="left", padx=(0, 15))
        ttk.Radiobutton(ma_frame, text="MA=$", variable=self.ma_var, value="$").pack(side="left")

        # Questionnaire File
        ttk.Label(input_container, text="แบบสอบถาม:", style='SubHeader.TLabel').grid(row=2, column=0, sticky="e", padx=(0, 10), pady=(10, 0))
        self.btn_load_q = ttk.Button(input_container, text="เลือกไฟล์แบบสอบถาม (Word/Excel/Text)", command=self.load_questionnaire_files, state="disabled", style='Primary.TButton')
        self.btn_load_q.grid(row=2, column=1, sticky="ew", pady=(10, 0))
        self.label_q_files = ttk.Label(input_container, text="ยังไม่ได้เลือกไฟล์", foreground="grey", style='Status.TLabel', wraplength=400)
        self.label_q_files.grid(row=2, column=2, sticky="ew", padx=(15, 0), pady=(10, 0))

        # Optional Export Button
        self.btn_export_excel = ttk.Button(input_container, text="กด SaveFile Var+Value เป็น Excel!!", command=self.export_to_excel, state="disabled")
        self.btn_export_excel.grid(row=3, column=1, sticky="ew", pady=(15, 5))
        
        # --- 2. ส่วนสร้าง Prompt ---
        ttk.Label(main_frame, text="2. สร้าง Prompt สำหรับ AI", style='Header.TLabel').pack(fill="x", anchor="w", pady=(0, 5))
        prompt_container = ttk.Frame(main_frame, padding=20, style='Content.TFrame')
        prompt_container.pack(fill="x", pady=(0, 20))
        prompt_container.columnconfigure((0, 1, 2), weight=1)

        self.btn_prompt_jod = ttk.Button(prompt_container, text="สร้าง Prompt โจทย์ (Var)", command=self.generate_prompt_jod, state="disabled", style='Success.TButton')
        self.btn_prompt_jod.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        self.btn_prompt_code = ttk.Button(prompt_container, text="สร้าง Prompt Code (Value)", command=self.generate_prompt_code, state="disabled", style='Primary.TButton')
        self.btn_prompt_code.grid(row=0, column=1, sticky="ew", padx=5)
        self.btn_open_web = ttk.Button(prompt_container, text="เปิด AI Studio", command=self.open_ai_studio, style='Accent.TButton')
        self.btn_open_web.grid(row=0, column=2, sticky="ew", padx=(5, 0))

        # --- 3. ส่วนของผลลัพธ์ และ บันทึก ---
        ttk.Label(main_frame, text="3. ผลลัพธ์ และการบันทึก", style='Header.TLabel').pack(fill="x", anchor="w", pady=(0, 5))
        output_container = ttk.Frame(main_frame, padding=20, style='Content.TFrame')
        output_container.pack(fill="both", expand=True)
        output_container.columnconfigure(0, weight=1)
        output_container.rowconfigure(1, weight=1) # ให้ ScrolledText ขยายตามแนวตั้ง

        ttk.Label(output_container, text="Prompt ที่สร้างขึ้นในกล่องนี้กด คัดลอก Prompt ไปแปะใน AI:", style='SubHeader.TLabel').grid(row=0, column=0, sticky="w", pady=(0, 10))
        
        self.output_text = scrolledtext.ScrolledText(
            output_container, height=10, 
            font=("Consolas", 11), wrap=tk.WORD, 
            relief="flat", borderwidth=0,
            bg=self.INPUT_BG, fg=self.INPUT_FG,
            insertbackground=self.TEXT_HEADER_COLOR
        )
        self.output_text.grid(row=1, column=0, sticky="nsew", pady=(0, 15))
        
        # --- Final action buttons ---
        final_action_frame = ttk.Frame(output_container, style='Content.TFrame')
        final_action_frame.grid(row=2, column=0, sticky="ew")
        final_action_frame.columnconfigure((0, 1, 2), weight=1)
        
        btn_copy = ttk.Button(final_action_frame, text="คัดลอก Prompt", command=self.copy_to_clipboard, style='Destructive.TButton')
        btn_copy.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        self.btn_load_ai_result = ttk.Button(final_action_frame, text="โหลดไฟล์ Value Excel ที่จะแปะลง SPSS", command=self.load_ai_result_excel, state="disabled")
        self.btn_load_ai_result.grid(row=0, column=1, sticky="ew", padx=5)
        self.btn_save_spss = ttk.Button(final_action_frame, text="บันทึกภาษา ENG ลง SPSS (.sav)", command=self.save_labels_to_spss, state="disabled", style='Purple.TButton')
        self.btn_save_spss.grid(row=0, column=2, sticky="ew", padx=(5, 0))


    # --- ฟังก์ชันจัดการ SPSS (ไม่มีการเปลี่ยนแปลง) ---
    def _process_spss_data(self, spss_file_path, delimiter):
        try:
            df, meta = pyreadstat.read_sav(spss_file_path)
            
            label_data = []
            for var_name, var_labels in meta.variable_value_labels.items():
                if delimiter in var_name and not var_name.endswith(f"{delimiter}1"):
                    continue
                for value, label in var_labels.items(): 
                    label_data.append({'Variable': var_name, 'Value': value, 'Label': label})
            
            labels_df = pd.DataFrame(label_data)
            if not labels_df.empty:
                labels_df['Label_EN'] = pd.NA

            filtered_column_names = [name for name in meta.column_names if delimiter not in name or name.endswith(f"{delimiter}1")]
            filtered_column_labels = [label for name, label in zip(meta.column_names, meta.column_labels) if delimiter not in name or name.endswith(f"{delimiter}1")]

            labels_dd = pd.DataFrame({
                'Name': filtered_column_names,
                'VAR_THA': filtered_column_labels
            })
            labels_dd['VAR_ENG'] = pd.NA
            
            return df, meta, labels_dd, labels_df
        except Exception as e:
            messagebox.showerror("SPSS Error", f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ SPSS:\n{e}")
            return None, None, None, None
        
    def load_spss_file(self):
        spss_path = filedialog.askopenfilename(title="เลือกไฟล์ SPSS", filetypes=(("SPSS Files", "*.sav"), ("All files", "*.*")))
        if not spss_path: return

        delimiter = self.ma_var.get()
        df, meta, df_var, df_value = self._process_spss_data(spss_path, delimiter)

        if df is not None:
            self.spss_filepath = spss_path
            self.original_spss_df = df
            self.spss_meta = meta
            self.spss_df_var = df_var
            self.spss_df_value = df_value

            filename = os.path.basename(spss_path)
            status_msg = (f"{filename}\n"
                          f"Var Labels: {len(df_var)}, Value Labels: {len(df_value)}")
            self.label_spss_file.config(text=status_msg, foreground=self.TEXT_COLOR)

            self.btn_export_excel.config(state="normal")
            self.btn_load_q.config(state="normal")
            self.btn_load_ai_result.config(state="normal")
            messagebox.showinfo("โหลด SPSS สำเร็จ", f"โหลดข้อมูลจาก {filename} เรียบร้อย\nขั้นตอนต่อไป: เลือกไฟล์แบบสอบถาม")
        else:
            self.label_spss_file.config(text="เกิดข้อผิดพลาดในการโหลดไฟล์", foreground=self.DESTRUCTIVE_COLOR)

    def export_to_excel(self):
        if self.spss_df_var is None or self.spss_df_value is None:
            messagebox.showerror("ไม่มีข้อมูล", "กรุณาโหลดไฟล์ SPSS ก่อนครับ")
            return

        excel_path = filedialog.asksaveasfilename(
            parent=self.root,
            title='Save Excel File',
            defaultextension=".xlsx",
            filetypes=[("Excel Files", "*.xlsx"), ("All files", "*.*")]
        )
        if not excel_path:
            return

        try:
            with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
                self.spss_df_var.to_excel(writer, index=False, sheet_name='Var')
                self.spss_df_value.to_excel(writer, index=False, sheet_name='Value')
            messagebox.showinfo("บันทึกสำเร็จ", f"ไฟล์ Excel ถูกบันทึกที่:\n{excel_path}")
        except PermissionError:
            messagebox.showerror(
                "ไม่สามารถเขียนไฟล์ได้",
                "ไฟล์ปลายทางอาจถูกเปิดอยู่ใน Excel หรือไม่มีสิทธิ์เขียน โปรดปิดไฟล์แล้วลองใหม่"
            )
        except Exception as e:
            messagebox.showerror("เกิดข้อผิดพลาด", f"ไม่สามารถบันทึกไฟล์ Excel ได้:\n{e}")

    def load_ai_result_excel(self):
        filepath = filedialog.askopenfilename(
            title="เลือกไฟล์ Excel ผลลัพธ์จาก AI",
            filetypes=(("Excel files", "*.xlsx *.xls"), ("All files", "*.*"))
        )
        if not filepath: return

        try:
            xls = pd.ExcelFile(filepath)
            loaded_something = False
            
            self.ai_result_var_df = None
            self.ai_result_value_df = None

            if 'Var' in xls.sheet_names:
                self.ai_result_var_df = pd.read_excel(xls, sheet_name='Var')
                loaded_something = True
                messagebox.showinfo("โหลดสำเร็จ", f"โหลดข้อมูลจากชีท 'Var' เรียบร้อย\nไฟล์: {os.path.basename(filepath)}")

            if 'Value' in xls.sheet_names:
                self.ai_result_value_df = pd.read_excel(xls, sheet_name='Value')
                loaded_something = True
                messagebox.showinfo("โหลดสำเร็จ", f"โหลดข้อมูลจากชีท 'Value' เรียบร้อย\nไฟล์: {os.path.basename(filepath)}")

            if loaded_something:
                self.btn_save_spss.config(state="normal")
                messagebox.showinfo("พร้อมบันทึก", "ข้อมูลจาก AI ถูกโหลดแล้ว\nขั้นตอนต่อไป: กด 'บันทึกผลลง SPSS'")
            else:
                messagebox.showwarning("ไม่พบข้อมูล", "ไม่พบชีทชื่อ 'Var' หรือ 'Value' ในไฟล์ Excel ที่เลือก")

        except Exception as e:
            messagebox.showerror("เกิดข้อผิดพลาด", f"ไม่สามารถอ่านไฟล์ Excel ได้:\n{e}")
            
    def generate_prompt_jod(self):
        if not self.questionnaire_data:
            messagebox.showwarning("ข้อมูลไม่ครบ", "กรุณาเลือก 'ไฟล์แบบสอบถาม' ก่อนครับ")
            return
        
        var_data_text = self.spss_df_var.to_csv(sep='\t', index=False, header=True)
        combined_data = f"{self.questionnaire_data}\n\n{var_data_text}"
        final_prompt = PROMPT_JOD.replace("[ *** วางข้อมูล Word และ Excel *** ]", combined_data)
        
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, final_prompt)
        self.last_prompt_type = 'jod'
        self.btn_save_spss.config(state="normal")
        messagebox.showinfo("สำเร็จ", "สร้าง 'Prompt โจทย์' (Var) เรียบร้อย!\nคัดลอก Prompt นี้ไปใช้กับ AI ได้เลย")

    def generate_prompt_code(self):
        if not self.questionnaire_data:
            messagebox.showwarning("ข้อมูลไม่ครบ", "กรุณาเลือก 'ไฟล์แบบสอบถาม' ก่อนครับ")
            return
            
        value_df_for_prompt = self.spss_df_value.rename(columns={'Label': 'Label_Th'})
        value_data_text = value_df_for_prompt[['Variable', 'Value', 'Label_Th']].to_csv(sep='\t', index=False, header=True)

        combined_data = f"{self.questionnaire_data}\n\n{value_data_text}"
        final_prompt = PROMPT_CODE.replace("[ *** วางข้อมูล Word และ Excel *** ]", combined_data)
        
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, final_prompt)
        self.last_prompt_type = 'code'
        self.btn_save_spss.config(state="normal")
        messagebox.showinfo("สำเร็จ", "สร้าง 'Prompt Code' (Value) เรียบร้อย!\nคัดลอก Prompt นี้ไปใช้กับ AI ได้เลย")

    # --- ส่วนที่เหลือของฟังก์ชัน (save_labels_to_spss, load_questionnaire_files, etc.) ไม่มีการเปลี่ยนแปลง ---
    def save_labels_to_spss(self):
        try:
            import pyspssio
        except Exception:
            messagebox.showerror(
                "ไลบรารีไม่พร้อม",
                "ไม่พบโมดูล 'pyspssio'\n\nติดตั้งด้วยคำสั่ง:\n  python -m pip install pyspssio"
            )
            return

        import io, pandas as pd, tkinter as tk
        update_var_df = None
        update_value_df = None

        def find_col(df, names):
            for name in names:
                for col in df.columns:
                    if col.lower() == name.lower():
                        return col
            return None

        if self.ai_result_var_df is not None:
            update_var_df = self.ai_result_var_df
        if self.ai_result_value_df is not None:
            update_value_df = self.ai_result_value_df

        ai_output_text = self.output_text.get("1.0", tk.END).strip()
        if update_var_df is None and update_value_df is None and ai_output_text:
            try:
                df_txt = pd.read_csv(io.StringIO(ai_output_text), sep="\t")
                if self.last_prompt_type == "jod":
                    update_var_df = df_txt
                elif self.last_prompt_type == "code":
                    update_value_df = df_txt
            except Exception as e:
                messagebox.showerror("Parse Error", f"อ่านข้อมูลจากกล่องข้อความไม่สำเร็จ:\n{e}")
                return
        if update_var_df is None and update_value_df is None:
            messagebox.showerror("ไม่มีข้อมูล", "กรุณาโหลดไฟล์ผลลัพธ์จาก AI หรือวางข้อมูลก่อน")
            return

        delimiter = self.ma_var.get()
        meta = self.spss_meta
        try:
            original_columns = meta.column_names
            df_to_save = self.original_spss_df[original_columns].copy()
        except KeyError as e:
            messagebox.showerror("เกิดข้อผิดพลาด", f"ไม่พบคอลัมน์: {e}")
            return

        new_var_labels_list = meta.column_labels[:]
        new_value_labels = meta.variable_value_labels.copy()

        if update_value_df is not None:
            lbl_en_col = find_col(update_value_df, ["Label_EN", "Label_En", "label_en"])
            if not lbl_en_col:
                messagebox.showerror("ไม่พบคอลัมน์", "ไม่พบ 'Label_EN' ในชีท Value")
                return
            update_value_df["Value"] = pd.to_numeric(update_value_df["Value"], errors="coerce")
            for _, r in update_value_df.iterrows():
                var, val, en = r["Variable"], r["Value"], str(r.get(lbl_en_col, "") or "")
                if pd.isna(val) or var not in df_to_save.columns: continue
                if en == "" or en.lower() in ("nan", "none"):
                    if var in new_value_labels and val in new_value_labels[var]:
                        del new_value_labels[var][val]
                else:
                    new_value_labels.setdefault(var, {})
                    try: vv = int(val)
                    except Exception: vv = float(val)
                    new_value_labels[var][vv] = en

        vars_with_ma = [v for v in df_to_save.columns if delimiter in v]
        for base in {v.split(delimiter)[0] for v in vars_with_ma}:
            key = f"{base}{delimiter}1"
            if key in new_value_labels:
                src = new_value_labels[key]
                for v in vars_with_ma:
                    if v.startswith(base) and v != key and v in df_to_save.columns:
                        new_value_labels[v] = src

        if update_var_df is not None:
            eng_col = find_col(update_var_df, ["VAR_ENG", "Var_Eng", "var_eng"])
            if not eng_col:
                messagebox.showerror("ไม่พบคอลัมน์", "ไม่พบ 'VAR_ENG' ในชีท Var")
                return
            for _, r in update_var_df.iterrows():
                name, eng = r["Name"], r.get(eng_col, None)
                if pd.notna(eng) and name in original_columns:
                    idx = original_columns.index(name)
                    new_var_labels_list[idx] = str(eng)

        for base in {v.split(delimiter)[0] for v in df_to_save.columns if delimiter in v}:
            key = f"{base}{delimiter}1"
            if key in original_columns:
                idx = original_columns.index(key)
                lbl = new_var_labels_list[idx]
                if pd.notna(lbl) and lbl != "":
                    for v in [vv for vv in df_to_save.columns if vv.startswith(f"{base}{delimiter}")]:
                        new_var_labels_list[original_columns.index(v)] = lbl

        var_labels = {c: ("" if pd.isna(l) else str(l)) for c, l in zip(original_columns, new_var_labels_list)}
        meta_out = {"var_labels": var_labels, "var_value_labels": new_value_labels}

        new_file = filedialog.asksaveasfilename(
            parent=self.root, title="บันทึกไฟล์ SPSS ใหม่ (pyspssio)",
            defaultextension=".sav", filetypes=[("SPSS Files", "*.sav"), ("All Files", "*.*")]
        )
        if not new_file: return

        try:
            pyspssio.write_sav(new_file, df_to_save, metadata=meta_out)
            messagebox.showinfo("บันทึกสำเร็จ", f"ไฟล์ถูกบันทึกเรียบร้อยแล้ว\n\n{new_file}")
            self.ai_result_var_df = None
            self.ai_result_value_df = None
        except Exception as e:
            messagebox.showerror("บันทึกไม่สำเร็จ", f"เขียนไฟล์ด้วย pyspssio ไม่สำเร็จ:\n{e}")
                 
    def load_questionnaire_files(self):
        filepaths = filedialog.askopenfilenames(
            title="เลือกไฟล์แบบสอบถาม",
            filetypes=(("All Supported Files", "*.txt *.docx *.xlsx *.xls"),
                       ("Text files", "*.txt"), ("Word documents", "*.docx"),
                       ("Excel workbooks", "*.xlsx *.xls"), ("All files", "*.*"))
        )
        if filepaths:
            all_contents = [self.read_general_file_content(path) for path in filepaths]
            self.questionnaire_data = "\n\n--- End of File ---\n\n".join(all_contents)
            filenames = [os.path.basename(path) for path in filepaths]
            self.label_q_files.config(text="- " + "\n- ".join(filenames), foreground=self.TEXT_COLOR)
            
            self.btn_prompt_jod.config(state="normal")
            self.btn_prompt_code.config(state="normal")
            messagebox.showinfo("โหลดสำเร็จ", f"โหลดข้อมูลจาก {len(filenames)} ไฟล์เรียบร้อย\nขั้นตอนต่อไป: สร้าง Prompt ที่ต้องการ")
    
    def read_general_file_content(self, filepath):
        _, extension = os.path.splitext(filepath)
        content = ""
        try:
            if extension == '.docx':
                if not LIBS_INSTALLED: return "[Error: ไม่สามารถอ่าน .docx ได้]"
                doc = docx.Document(filepath)
                full_text = [para.text for para in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                content = '\n'.join(full_text)
            elif extension in ['.xlsx', '.xls']:
                if not LIBS_INSTALLED: return "[Error: ไม่สามารถอ่าน Excel ได้]"
                workbook = openpyxl.load_workbook(filepath, data_only=True)
                full_text = []
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_text = []
                    for row in sheet.iter_rows():
                        row_text = [str(cell.value) if cell.value is not None else "" for cell in row]
                        sheet_text.append("\t".join(row_text))
                    full_text.append('\n\n'.join(sheet_text))
                content = '\n\n'.join(full_text)
            else:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
        except Exception as e:
            content = f"[เกิดข้อผิดพลาดในการอ่านไฟล์ {os.path.basename(filepath)}: {e}]"
        return content

    def open_ai_studio(self):
        url = "https://aistudio.google.com/prompts/new_chat?lfhs=2"
        try:
            webbrowser.open_new_tab(url)
        except Exception as e:
            messagebox.showerror("เกิดข้อผิดพลาด", f"ไม่สามารถเปิดเบราว์เซอร์ได้\n\nError: {e}")
            
    def copy_to_clipboard(self):
        text_to_copy = self.output_text.get("1.0", tk.END).strip()
        if text_to_copy:
            pyperclip.copy(text_to_copy)
            messagebox.showinfo("คัดลอกสำเร็จ", "คัดลอกข้อความลงในคลิปบอร์ดแล้ว!")
        else:
            messagebox.showwarning("ไม่มีข้อความ", "ไม่มีข้อความในช่องผลลัพธ์ให้คัดลอก")


    
    
# <<< START OF CHANGES >>>
# --- ฟังก์ชัน Entry Point ใหม่ (สำหรับให้ Launcher เรียก) ---
def run_this_app(working_dir=None): # ชื่อฟังก์ชันนี้จะถูกใช้ใน Launcher
    """
    ฟังก์ชันหลักสำหรับสร้างและรัน QuotaSamplerApp.
    """
    print(f"--- QUOTA_SAMPLER_INFO: Starting 'QuotaSamplerApp' via run_this_app() ---")
    try:
    # --- ส่วนที่ใช้รันโปรแกรม ---
    #if __name__ == "__main__":
        root = tk.Tk()
        app = App(root)
        root.mainloop()
        print(f"--- QUOTA_SAMPLER_INFO: QuotaSamplerApp mainloop finished. ---")

    except Exception as e:
        # ดักจับ Error ที่อาจเกิดขึ้นระหว่างการสร้างหรือรัน App
        print(f"QUOTA_SAMPLER_ERROR: An error occurred during QuotaSamplerApp execution: {e}")
        # แสดง Popup ถ้ามีปัญหา
        if 'root' not in locals() or not root.winfo_exists(): # สร้าง root ชั่วคราวถ้ายังไม่มี
            root_temp = tk.Tk()
            root_temp.withdraw()
            messagebox.showerror("Application Error (Quota Sampler)",
                               f"An unexpected error occurred:\n{e}", parent=root_temp)
            root_temp.destroy()
        else:
            messagebox.showerror("Application Error (Quota Sampler)",
                               f"An unexpected error occurred:\n{e}", parent=root) # ใช้ root ที่มีอยู่ถ้าเป็นไปได้
        sys.exit(f"Error running QuotaSamplerApp: {e}") # อาจจะ exit หรือไม่ก็ได้ ขึ้นกับการออกแบบ


# --- ส่วน Run Application เมื่อรันไฟล์นี้โดยตรง (สำหรับ Test) ---
if __name__ == "__main__":
    print("--- Running QuotaSamplerApp.py directly for testing ---")
    # (ถ้ามีการตั้งค่า DPI ด้านบน มันจะทำงานอัตโนมัติ)

    # เรียกฟังก์ชัน Entry Point ที่เราสร้างขึ้น
    run_this_app()

    print("--- Finished direct execution of QuotaSamplerApp.py ---")
# <<< END OF CHANGES >>>