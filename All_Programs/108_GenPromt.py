import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog
import os
import pyperclip
import webbrowser # <-- [1] เพิ่ม import

# --- ส่วนของการอ่านไฟล์ ---
try:
    import docx
    import openpyxl
    LIBS_INSTALLED = True
except ImportError:
    LIBS_INSTALLED = False

# --- แม่แบบ Prompt (Templates) ---
# ... (วางแม่แบบ PROMPT_JOD และ PROMPT_CODE ทั้งหมดไว้ที่นี่ เหมือนเดิม) ...
PROMPT_JOD = """คุยไทยนะ
จากข้อมูลต่อไปนี้:

[ *** วางข้อมูล Word และ Excel *** ]

ให้ดำเนินการดังนี้:

1.  **ระบุคู่ข้อมูล:** สำหรับแต่ละแถวในตารางข้อมูลสุดท้าย (`Name`, `VAR_THA`):
    *   ค้นหาข้อความภาษาไทย (`VAR_THA`) ที่ตรงกันในส่วนข้อมูลต้นฉบับ
    *   ถ้าใน (`VAR_THA`) มีคำว่า Hide ไม่ต้องสนใจให้ข้ามบรรทัดนั้นไปแต่ให้คงตำแหน่งข้อความใน (`VAR_THA`) ไว้เหมือนเดิม
    *   มองหาข้อความภาษาอังกฤษ**ที่ปรากฏคู่กันโดยตรง**กับ `VAR_THA` นั้นๆ ในข้อมูลต้นฉบับ (เช่น อยู่ในคอลัมน์ติดกัน, อยู่ในบรรทัดเดียวกันในตำแหน่งที่สอดคล้องกัน, อยู่ในลำดับถัดกัน (ก่อนหน้า / หลัง) ภายในกลุ่มเดียวกัน, หรือมีการระบุว่าเป็นคำแปลอย่างชัดเจน)

2.  **สร้างตารางผลลัพธ์:** สร้างตารางที่มี 3 คอลัมน์ ตามลำดับดังนี้: `Name`, `VAR_THA`, `VAR_ENG`

3.  **เติมข้อมูล:**
    *   นำข้อมูล `Name` และ `VAR_THA` จากตารางสุดท้ายมาใส่ในตารางผลลัพธ์
    *   **เฉพาะกรณีที่** พบข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**ตามเงื่อนไขในข้อ 1 เท่านั้น:
        *   **การสร้างส่วนนำหน้า (prefix) สำหรับ VAR_ENG:**
            *   ให้สกัดหมายเลขคำถาม (เช่น S1, Q1, MQ1, Q2.1) จากส่วนเริ่มต้นของข้อมูลในคอลัมน์ `Name` (เช่น ถ้า `Name` คือ `s1` ให้ใช้ "S1", ถ้า `Name` คือ `q1_1_O1` ให้ใช้ "Q1.1", ถ้า `Name` คือ `mq5` ให้ใช้ "MQ5") หรือจากส่วนเริ่มต้นของ `VAR_THA` หากมีการระบุหมายเลขคำถามไว้อย่างชัดเจน (เช่น "S1 กรุณาระบุ...", "Q1.1 ...")
            *   นำหมายเลขคำถามที่สกัดได้ตามด้วยเครื่องหมายจุด (.) เพื่อสร้างเป็นส่วนนำหน้า (prefix)
        *   นำส่วนนำหน้าที่สร้างขึ้นนี้ (หมายเลขคำถาม + จุด) ต่อด้วยข้อความภาษาอังกฤษที่จับคู่ได้ มาเติมในคอลัมน์ `VAR_ENG` ของแถวที่สอดคล้องกัน
        *   ตัวอย่างเช่น:
            *   ถ้า `Name` คือ `mq5`, `VAR_THA` คือ "MQ5 คุณรู้สึกอยากซื้อสินค้า...", และข้อความอังกฤษที่คู่กันคือ "To what extent would you want to purchase the product..." ผลลัพธ์ใน `VAR_ENG` ควรเป็น "MQ5.To what extent would you want to purchase the product..."
            *   ถ้า `Name` คือ `I_1_q1`, `VAR_THA` คือ "(บิสกิต (รวมถึง ชนิดแท่ง)) Q1 กรุณาระบุความถี่...", และข้อความอังกฤษที่คู่กันคือ "Biscuit" ผลลัพธ์ใน `VAR_ENG` ควรเป็น "Q1.Biscuit" (โดย "Q1." มาจากการตีความ `I_1_q1` หรือ "Q1" ใน `VAR_THA`)
            *   ถ้า `Name` คือ `s1`, `VAR_THA` คือ "S1 กรุณาระบุเพศ...", และข้อความอังกฤษที่คู่กันคือ "Please indicate your gender..." ผลลัพธ์ใน `VAR_ENG` ควรเป็น "S1.Please indicate your gender..."
    *   **หากไม่พบ**ข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**กับ `VAR_THA` ในตำแหน่งที่สอดคล้องกันในข้อมูลต้นฉบับ ให้**เว้นว่าง**คอลัมน์ `VAR_ENG` สำหรับแถวนั้นไว้ **ห้าม**นำข้อมูลภาษาอังกฤษจากส่วนอื่นของเอกสารที่ไม่ใช่คู่โดยตรงมาเติม หรือทำการแปลเอง
    *   ตัวอย่างรูปแบบที่ต้องการสำหรับแถวที่มีการเติมเลขข้อ (เป็นการยืนยันรูปแบบ ไม่ใช่ให้ทำเฉพาะ P16):
        `P16 คุณรู้สึกไม่ชอบผลิตภัณฑ์ที่ได้ดื่มไปเมื่อสักครู่นี้ในแง่ใดบ้าง กรุณาตอบทุกข้อที่ไม่ชอบ P16.Regarding the product you drank earlier, please select all the aspects you disliked from the following options.`

4.  **จัดรูปแบบ VAR_ENG:**
    *   สำหรับ `VAR_ENG` ที่มีการเติมข้อมูล: ตรวจสอบให้แน่ใจว่าข้อความ**ไม่มีช่องว่างนำหน้า (leading space)** และ **ไม่มีช่องว่างตามหลัง (trailing space)** (ยกเว้นช่องว่างที่อาจมีระหว่าง Prefix ที่สร้างกับตัวข้อความภาษาอังกฤษ ถ้าการจับคู่เดิมมีช่องว่างนั้นอยู่)
    *   **คง**ช่องว่างที่อยู่**ภายใน**ข้อความภาษาอังกฤษไว้ตามเดิม

5.  **จัดรูปแบบตาราง:**
    *   แสดงผลลัพธ์เป็น **ข้อความธรรมดา (Plain Text)**
    *   ใช้ **ตัวคั่นแท็บ (Tab character)** เพียงตัวเดียวในการแยกระหว่างข้อมูลในแต่ละคอลัมน์
    *   **ห้าม**รวมเซลล์ (No Merging)
    ส่งกลับมาเป็น Table Markdown ที่มี 3 คอลัมน์: `Name`, `VAR_THA`, `VAR_ENG`
    จัดให้อยู่ในรูปแบบ Codebox เป็นตาราง Excel
"""
PROMPT_CODE = """คุยไทยนะ
จากข้อมูลต่อไปนี้:

[ *** วางข้อมูล Word และ Excel *** ]

ให้ดำเนินการดังนี้:

1.  **ระบุคู่ข้อมูล:** สำหรับแต่ละแถวในตารางข้อมูลสุดท้าย (`Variable`, `Value`, `Label_Th`):
    *   ค้นหาข้อความภาษาไทย (`Label_Th`) ที่ตรงกันในส่วนข้อมูลต้นฉบับ
    *   ถ้าใน (`Label_Th`) มีคำว่า Hide ไม่ต้องสนใจให้ข้ามบรรทัดนั้นไปแต่ให้คงตำแหน่งข้อความใน (`Label_Th`) ไว้เหมือนเดิม
    *   มองหาข้อความภาษาอังกฤษ**ที่ปรากฏคู่กันโดยตรง**กับ `Label_Th` นั้นๆ ในข้อมูลต้นฉบับ (เช่น อยู่ในคอลัมน์ติดกัน, อยู่ในบรรทัดเดียวกันในตำแหน่งที่สอดคล้องกัน อยู่ในลำดับถัดกัน (ก่อนหน้า / หลัง) ภายในกลุ่มเดียวกัน ,หรือมีการระบุว่าเป็นคำแปลอย่างชัดเจน)
2.  **สร้างตารางผลลัพธ์:** สร้างตารางที่มี 4 คอลัมน์ ตามลำดับดังนี้: `Variable`, `Value`, `Label_Th`, `Label_En`
3.  **เติมข้อมูล:**
    *   นำข้อมูล `Variable`, `Value`, และ `Label_Th` จากตารางสุดท้ายมาใส่ในตารางผลลัพธ์
    *   **เฉพาะกรณีที่** พบข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**ตามเงื่อนไขในข้อ 1 เท่านั้น ให้นำข้อความภาษาอังกฤษนั้นมาเติมในคอลัมน์ `Label_En` ของแถวที่สอดคล้องกัน
    *   **หากไม่พบ**ข้อความภาษาอังกฤษที่**จับคู่กันโดยตรง**กับ `Label_Th` ในตำแหน่งที่สอดคล้องกันในข้อมูลต้นฉบับ ให้**เว้นว่าง**คอลัมน์ `Label_En` สำหรับแถวนั้นไว้ **ห้าม**นำข้อมูลภาษาอังกฤษจากส่วนอื่นของเอกสารที่ไม่ใช่คู่โดยตรงมาเติม หรือทำการแปลเอง
4.  **จัดรูปแบบ Label_En:**
    *   สำหรับ `Label_En` ที่มีการเติมข้อมูล: ตรวจสอบให้แน่ใจว่าข้อความ**ไม่มีช่องว่างนำหน้า (leading space)** และ **ไม่มีช่องว่างตามหลัง (trailing space)**
    *   **คง**ช่องว่างที่อยู่**ภายใน**ข้อความภาษาอังกฤษไว้ตามเดิม
	ถ้ามี (R1) (R2) หน้า Code ให้เอามาด้วย เช่น
	Variable	Value	Label	Label_EN
	s16a_O1	1	(R1) โซฟี ขอบปกป้อง ผิวสัมผัสแห้ง สลิม มีปีก	(R1) Sofy Side Gather slim (wing)
	s16a_O1	2	(R2) โซฟี ขอบปกป้อง ผิวสัมผัสนุ่ม สลิม มีปีก	(R2) Sofy Side Gather slim Non-mesh (wing)
5.  **จัดรูปแบบตาราง:**
    *   แสดงผลลัพธ์เป็น **ข้อความธรรมดา (Plain Text)**
    *   ใช้ **ตัวคั่นแท็บ (Tab character)** เพียงตัวเดียวในการแยกระหว่างข้อมูลในแต่ละคอลัมน์
    *   **ห้าม**รวมเซลล์ (No Merging)
    ส่งกลับมาเป็น Table Markdown ที่มี 4 คอลัมน์: `Variable`, `Value`, `Label_Th`, `Label_En`
    จัดให้อยู่ในรูปแบบ Codebox เป็นตาราง Excel
"""


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("โปรแกรมสร้าง Promt แปะ Eng v1")
        
        # --- โค้ดสำหรับจัดหน้าต่างให้อยู่กลางจอ ---
        window_width = 1000
        window_height = 750

        # ดึงขนาดของหน้าจอคอมพิวเตอร์
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # คำนวณหาจุดกึ่งกลางสำหรับแกน x และ y
        center_x = int(screen_width/2 - window_width / 2)
        center_y = int(screen_height/2 - window_height / 2)

        # ตั้งค่าตำแหน่งและขนาดของหน้าต่าง
        self.root.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')

        # ตัวแปรสำหรับเก็บข้อมูลที่อ่านจากไฟล์
        self.questionnaire_data = ""
        # เปลี่ยนเป็น dictionary เพื่อเก็บข้อมูลแยกชีท
        self.cat_program_data = {"var": "", "value": ""}
        self.cat_program_filepath = ""

        if not LIBS_INSTALLED:
            messagebox.showerror("ไลบรารีไม่ครบ", 
                                 "ยังไม่ได้ติดตั้ง python-docx หรือ openpyxl\n"
                                 "โปรแกรมจะสามารถอ่านได้เฉพาะไฟล์ .txt เท่านั้น\n\n"
                                 "กรุณาติดตั้งด้วยคำสั่ง:\npip install python-docx openpyxl")

        # --- ส่วนของ Input ---
        input_frame = tk.Frame(root)
        input_frame.pack(pady=10, padx=10, fill="x", anchor="n")
        input_frame.grid_columnconfigure(0, weight=1)
        input_frame.grid_columnconfigure(1, weight=1)

        # Input 1: แบบสอบถาม
        frame1 = tk.LabelFrame(input_frame, text=" 1. ไฟล์แบบสอบถาม ", font=("Tahoma", 10, "bold"), padx=10, pady=10)
        frame1.grid(row=0, column=0, sticky="nsew", padx=5)
        
        btn_load1 = tk.Button(frame1, text="เลือกไฟล์แบบสอบถาม (เลือกหลายไฟล์ได้)", command=self.load_questionnaire_files, font=("Tahoma", 10))
        btn_load1.pack(fill="x", pady=5)
        self.label_files1 = tk.Label(frame1, text="ยังไม่ได้เลือกไฟล์", justify="left", fg="grey", wraplength=450, height=5, anchor="nw")
        self.label_files1.pack(fill="x", pady=5)

        # Input 2: โปรแกรมแมว
        frame2 = tk.LabelFrame(input_frame, text=" 2. ไฟล์โปรแกรมแมว ", font=("Tahoma", 10, "bold"), padx=10, pady=10)
        frame2.grid(row=0, column=1, sticky="nsew", padx=5)

        btn_load2 = tk.Button(frame2, text="เลือกไฟล์โปรแกรมแมว (Excel)", command=self.load_cat_program_file, font=("Tahoma", 10))
        btn_load2.pack(fill="x", pady=5)
        self.label_files2 = tk.Label(frame2, text="ยังไม่ได้เลือกไฟล์", justify="left", fg="grey", wraplength=450, height=5, anchor="nw")
        self.label_files2.pack(fill="x", pady=5)
        
        # --- ส่วนของปุ่มคำสั่ง ---
        button_frame = tk.Frame(root)
        button_frame.pack(pady=5, padx=10, fill="x")
        
        self.btn_prompt_jod = tk.Button(button_frame, text="สร้าง Prompt โจทย์", command=self.generate_prompt_jod, font=("Tahoma", 11, "bold"), bg="#4CAF50", fg="white")
        self.btn_prompt_jod.pack(side="left", padx=10, pady=5, ipadx=10, ipady=5)

        self.btn_prompt_code = tk.Button(button_frame, text="สร้าง Prompt Code", command=self.generate_prompt_code, font=("Tahoma", 11, "bold"), bg="#2196F3", fg="white")
        self.btn_prompt_code.pack(side="left", padx=10, pady=5, ipadx=10, ipady=5)
        
        # --- [3][START] ปุ่มที่เพิ่มเข้ามา ---
        self.btn_open_web = tk.Button(button_frame, text="เปิด AI Studio", command=self.open_ai_studio, font=("Tahoma", 11, "bold"), bg="#FF9800", fg="white")
        self.btn_open_web.pack(side="left", padx=10, pady=5, ipadx=10, ipady=5)
        # --- [3][END] ปุ่มที่เพิ่มเข้ามา ---

        # --- ส่วนของผลลัพธ์ ---
        output_frame = tk.Frame(root)
        output_frame.pack(pady=10, padx=10, fill="both", expand=True)
        
        output_label = tk.Label(output_frame, text="ผลลัพธ์ (Prompt ที่สร้างเสร็จ):", font=("Tahoma", 10))
        output_label.pack(anchor="w")
        
        self.output_text = scrolledtext.ScrolledText(output_frame, height=15, font=("Tahoma", 10), wrap=tk.WORD)
        self.output_text.pack(fill="both", expand=True, pady=(0, 5))
        
        self.btn_copy = tk.Button(output_frame, text="คัดลอกผลลัพธ์ทั้งหมด", command=self.copy_to_clipboard, font=("Tahoma", 11, "bold"), bg="#f44336", fg="white")
        self.btn_copy.pack(pady=5, ipadx=10, ipady=5)

    def read_text_from_sheet(self, sheet):
        """อ่านข้อมูลทั้งหมดจากชีทที่กำหนดให้เป็น Plain Text"""
        full_text = []
        for row in sheet.iter_rows():
            row_text = [str(cell.value) if cell.value is not None else "" for cell in row]
            full_text.append("\t".join(row_text))
        return '\n'.join(full_text)

    def read_general_file_content(self, filepath):
        """ฟังก์ชันสำหรับอ่านไฟล์แบบสอบถาม (Word, Excel, Text)"""
        _, extension = os.path.splitext(filepath)
        content = ""
        try:
            if extension == '.docx':
                if not LIBS_INSTALLED: return "[Error: ไม่สามารถอ่าน .docx ได้]"
                doc = docx.Document(filepath)
                full_text = [para.text for para in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                content = '\n'.join(full_text)
            elif extension in ['.xlsx', '.xls']:
                if not LIBS_INSTALLED: return "[Error: ไม่สามารถอ่าน Excel ได้]"
                workbook = openpyxl.load_workbook(filepath, data_only=True)
                full_text = []
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    full_text.append(self.read_text_from_sheet(sheet))
                content = '\n\n'.join(full_text) # แยกแต่ละชีทด้วยบรรทัดว่าง
            else: # .txt หรือไฟล์อื่นๆ
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
        except Exception as e:
            content = f"[เกิดข้อผิดพลาดในการอ่านไฟล์ {os.path.basename(filepath)}: {e}]"
        return content

    def load_questionnaire_files(self):
        filepaths = filedialog.askopenfilenames(
            title="เลือกไฟล์แบบสอบถาม",
            filetypes=(("All Supported Files", "*.txt *.docx *.xlsx *.xls"),
                       ("Text files", "*.txt"),
                       ("Word documents", "*.docx"),
                       ("Excel workbooks", "*.xlsx *.xls"),
                       ("All files", "*.*"))
        )
        if filepaths:
            all_contents = [self.read_general_file_content(path) for path in filepaths]
            self.questionnaire_data = "\n\n--- End of File ---\n\n".join(all_contents)
            filenames = [os.path.basename(path) for path in filepaths]
            self.label_files1.config(text="ไฟล์ที่เลือก:\n- " + "\n- ".join(filenames), fg="black")
            messagebox.showinfo("โหลดสำเร็จ", f"โหลดข้อมูลจาก {len(filenames)} ไฟล์เรียบร้อย")

    def load_cat_program_file(self):
        """แก้ไขให้ดึงข้อมูลจากชีท 'Var' และ 'Value' โดยเฉพาะ"""
        filepath = filedialog.askopenfilename(
            title="เลือกไฟล์โปรแกรมแมว (Excel)",
            filetypes=(("Excel workbooks", "*.xlsx *.xls"), ("All files", "*.*"))
        )
        if filepath and LIBS_INSTALLED:
            try:
                workbook = openpyxl.load_workbook(filepath, data_only=True)
                sheet_names = [name.lower() for name in workbook.sheetnames]
                
                # รีเซ็ตข้อมูลเก่า
                self.cat_program_data = {"var": "", "value": ""}
                
                found_var = False
                found_value = False
                
                if "var" in sheet_names:
                    sheet_var = workbook[workbook.sheetnames[sheet_names.index("var")]]
                    self.cat_program_data["var"] = self.read_text_from_sheet(sheet_var)
                    found_var = True

                if "value" in sheet_names:
                    sheet_value = workbook[workbook.sheetnames[sheet_names.index("value")]]
                    self.cat_program_data["value"] = self.read_text_from_sheet(sheet_value)
                    found_value = True

                self.cat_program_filepath = filepath
                filename = os.path.basename(filepath)
                
                status_message = f"ไฟล์ที่เลือก:\n- {filename}\n\nสถานะ:\n"
                status_message += f"- พบชีท 'Var': {'ใช่' if found_var else 'ไม่'}\n"
                status_message += f"- พบชีท 'Value': {'ใช่' if found_value else 'ไม่'}"

                self.label_files2.config(text=status_message, fg="black")

                if not found_var and not found_value:
                    messagebox.showwarning("ไม่พบชีท", f"ไม่พบชีทชื่อ 'Var' หรือ 'Value' ในไฟล์ {filename}")
                else:
                    messagebox.showinfo("โหลดสำเร็จ", f"โหลดข้อมูลจาก '{filename}' เรียบร้อย")
            
            except Exception as e:
                messagebox.showerror("เกิดข้อผิดพลาด", f"ไม่สามารถอ่านไฟล์ Excel '{os.path.basename(filepath)}' ได้\n\nError: {e}")
        elif not LIBS_INSTALLED:
             messagebox.showerror("ไลบรารีไม่ครบ", "กรุณาติดตั้ง openpyxl เพื่ออ่านไฟล์ Excel")


    def generate_prompt_jod(self):
        """ปุ่มเขียว: แบบสอบถาม + ชีท Var"""
        if not self.questionnaire_data:
            messagebox.showwarning("ข้อมูลไม่ครบ", "กรุณาเลือก 'ไฟล์แบบสอบถาม' ก่อนครับ")
            return
        if not self.cat_program_data.get("var"):
            messagebox.showwarning("ข้อมูลไม่ครบ", "ไม่พบข้อมูลจากชีท 'Var' ในไฟล์โปรแกรมแมว\nกรุณาเลือกและตรวจสอบไฟล์อีกครั้ง")
            return
        
        combined_data = f"{self.questionnaire_data}\n\n{self.cat_program_data['var']}"
        final_prompt = PROMPT_JOD.replace("[ *** วางข้อมูล Word และ Excel *** ]", combined_data)
        
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, final_prompt)
        messagebox.showinfo("สำเร็จ", "สร้าง 'Prompt โจทย์' (แบบสอบถาม + Var) เรียบร้อยแล้ว!")

    def generate_prompt_code(self):
        """ปุ่มฟ้า: แบบสอบถาม + ชีท Value"""
        if not self.questionnaire_data:
            messagebox.showwarning("ข้อมูลไม่ครบ", "กรุณาเลือก 'ไฟล์แบบสอบถาม' ก่อนครับ")
            return
        if not self.cat_program_data.get("value"):
            messagebox.showwarning("ข้อมูลไม่ครบ", "ไม่พบข้อมูลจากชีท 'Value' ในไฟล์โปรแกรมแมว\nกรุณาเลือกและตรวจสอบไฟล์อีกครั้ง")
            return
            
        combined_data = f"{self.questionnaire_data}\n\n{self.cat_program_data['value']}"
        final_prompt = PROMPT_CODE.replace("[ *** วางข้อมูล Word และ Excel *** ]", combined_data)
        
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, final_prompt)
        messagebox.showinfo("สำเร็จ", "สร้าง 'Prompt Code' (แบบสอบถาม + Value) เรียบร้อยแล้ว!")
    
    # --- [2] เพิ่มฟังก์ชันสำหรับเปิดเว็บ ---
    def open_ai_studio(self):
        """เปิดหน้าเว็บ AI Studio ในเบราว์เซอร์"""
        url = "https://aistudio.google.com/prompts/new_chat?lfhs=2"
        try:
            webbrowser.open_new_tab(url)
        except Exception as e:
            messagebox.showerror("เกิดข้อผิดพลาด", f"ไม่สามารถเปิดเบราว์เซอร์ได้\n\nError: {e}")
            
    def copy_to_clipboard(self):
        text_to_copy = self.output_text.get("1.0", tk.END).strip()
        if text_to_copy:
            pyperclip.copy(text_to_copy)
            messagebox.showinfo("คัดลอกสำเร็จ", "คัดลอกข้อความลงในคลิปบอร์ดแล้ว!")
        else:
            messagebox.showwarning("ไม่มีข้อความ", "ไม่มีข้อความในช่องผลลัพธ์ให้คัดลอก")



# <<< START OF CHANGES >>>
# --- ฟังก์ชัน Entry Point ใหม่ (สำหรับให้ Launcher เรียก) ---
def run_this_app(working_dir=None): # ชื่อฟังก์ชันนี้จะถูกใช้ใน Launcher
    """
    ฟังก์ชันหลักสำหรับสร้างและรัน QuotaSamplerApp.
    """
    print(f"--- QUOTA_SAMPLER_INFO: Starting 'QuotaSamplerApp' via run_this_app() ---")
    try:
        # --- โค้ดที่ย้ายมาจาก if __name__ == "__main__": เดิมจะมาอยู่ที่นี่ ---
        # --- ส่วนที่ใช้รันโปรแกรม ---
        #if __name__ == "__main__":
        root = tk.Tk()
        app = App(root)
        root.mainloop()
        print(f"--- QUOTA_SAMPLER_INFO: QuotaSamplerApp mainloop finished. ---")

    except Exception as e:
        # ดักจับ Error ที่อาจเกิดขึ้นระหว่างการสร้างหรือรัน App
        print(f"QUOTA_SAMPLER_ERROR: An error occurred during QuotaSamplerApp execution: {e}")
        # แสดง Popup ถ้ามีปัญหา
        if 'root' not in locals() or not root.winfo_exists(): # สร้าง root ชั่วคราวถ้ายังไม่มี
            root_temp = tk.Tk()
            root_temp.withdraw()
            messagebox.showerror("Application Error (Quota Sampler)",
                               f"An unexpected error occurred:\n{e}", parent=root_temp)
            root_temp.destroy()
        else:
            messagebox.showerror("Application Error (Quota Sampler)",
                               f"An unexpected error occurred:\n{e}", parent=root) # ใช้ root ที่มีอยู่ถ้าเป็นไปได้
        sys.exit(f"Error running QuotaSamplerApp: {e}") # อาจจะ exit หรือไม่ก็ได้ ขึ้นกับการออกแบบ


# --- ส่วน Run Application เมื่อรันไฟล์นี้โดยตรง (สำหรับ Test) ---
if __name__ == "__main__":
    print("--- Running QuotaSamplerApp.py directly for testing ---")
    # (ถ้ามีการตั้งค่า DPI ด้านบน มันจะทำงานอัตโนมัติ)

    # เรียกฟังก์ชัน Entry Point ที่เราสร้างขึ้น
    run_this_app()

    print("--- Finished direct execution of QuotaSamplerApp.py ---")
# <<< END OF CHANGES >>>
